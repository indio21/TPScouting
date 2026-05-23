from __future__ import annotations

import sys
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
WORD_PATH = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_WORD_PATH = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagramas" / "export"


def norm(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_text(paragraph: Paragraph) -> str:
    return norm(paragraph.text)


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro parrafo exacto: {exact}")


def find_contains(doc: Document, fragment: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if fragment in paragraph_text(paragraph):
            return paragraph
    raise ValueError(f"No se encontro fragmento: {fragment}")


def find_caption(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Caption" and paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro caption exacto: {exact}")


def paragraph_after(doc: Document, cursor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    cursor._element.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def add_figure_after(
    doc: Document,
    cursor: Paragraph,
    image_name: str,
    caption: str,
    intro: str | None = None,
    width_inches: float = 6.35,
) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Caption" and paragraph_text(paragraph) == caption:
            return paragraph

    current = cursor
    if intro:
        current = paragraph_after(doc, current, intro)
    image_paragraph = paragraph_after(doc, current)
    image_paragraph.add_run().add_picture(str(DIAGRAM_DIR / image_name), width=Inches(width_inches))
    return paragraph_after(doc, image_paragraph, caption, "Caption")


def update_figure_list(doc: Document) -> None:
    list_heading = find_paragraph(doc, "LISTA DE FIGURAS")
    table_heading = find_paragraph(doc, "LISTA DE TABLAS")

    node = list_heading._element.getnext()
    while node is not None and node is not table_heading._element:
        next_node = node.getnext()
        text = norm(Paragraph(node, doc._body).text) if node.tag.endswith("}p") else ""
        if text.startswith("Figura "):
            node.getparent().remove(node)
        node = next_node

    entries = [
        "Figura 4-1. Diagrama de componentes de TPScouting.",
        "Figura 4-2. Diagrama de clases del modelo de dominio.",
        "Figura 5-1. Secuencia de prediccion de potencial de jugador.",
        "Figura 5-2. Secuencia de carga del panel general.",
        "Figura 5-3. Diagrama de despliegue en Render.",
        "Figura 6-1. Curva real de entrenamiento desde training_metadata.json.",
        "Figura 6-2. Login del MVP desplegado en Render.",
        "Figura 6-3. Panel general o mesa de scouting.",
        "Figura 6-4. Listado paginado de jugadores.",
        "Figura 6-5. Ficha de jugador con historial.",
        "Figura 6-6. Vista de predicción de potencial.",
        "Figura 6-7. Comparador múltiple de jugadores.",
    ]
    for entry in entries:
        # Insertar en orden estable dejando LISTA DE TABLAS donde estaba.
        paragraph = paragraph_after(doc, list_heading, entry)
        paragraph._element.getparent().remove(paragraph._element)
        table_heading._element.addprevious(paragraph._element)


def main() -> None:
    doc = Document(WORD_PATH)

    update_figure_list(doc)

    add_figure_after(
        doc,
        find_contains(doc, "Interacción del usuario: Los usuarios pueden consultar datos de jugadores"),
        "03_componentes.png",
        "Figura 4-1. Diagrama de componentes de TPScouting.",
        "La Figura 4-1 resume los componentes principales de la solucion implementada: cliente web, aplicacion Flask, blueprints funcionales, servicios auxiliares, persistencia y modulo de Machine Learning.",
    )
    add_figure_after(
        doc,
        find_caption(doc, "Tabla 4-7. Entidades principales del modelo de datos."),
        "05_clases.png",
        "Figura 4-2. Diagrama de clases del modelo de dominio.",
        "La Figura 4-2 muestra las clases principales del modelo de dominio persistido con SQLAlchemy y sus relaciones mas relevantes.",
    )
    add_figure_after(
        doc,
        find_contains(doc, "Flujo técnico de inferencia: el usuario ingresa a la vista de predicción"),
        "01_secuencia_prediccion.png",
        "Figura 5-1. Secuencia de prediccion de potencial de jugador.",
        "La Figura 5-1 detalla el flujo de inferencia cuando un usuario consulta la proyeccion de potencial de un jugador.",
    )
    add_figure_after(
        doc,
        find_contains(doc, "El frontend utiliza Jinja2, Bootstrap, CSS propio y Chart.js"),
        "02_secuencia_dashboard.png",
        "Figura 5-2. Secuencia de carga del panel general.",
        "La Figura 5-2 describe la carga del panel general, incluyendo validacion de sesion, cache in-memory, consultas SQLAlchemy y renderizado de la vista.",
    )
    add_figure_after(
        doc,
        find_contains(doc, "En Render se evita SQLite como almacenamiento productivo"),
        "04_despliegue.png",
        "Figura 5-3. Diagrama de despliegue en Render.",
        "La Figura 5-3 representa el despliegue verificado del MVP en Render con Gunicorn, Flask, PostgreSQL y artefactos de Machine Learning.",
    )

    doc.save(WORD_PATH)
    copy2(WORD_PATH, REPO_WORD_PATH)
    print(f"Diagramas insertados en: {WORD_PATH}")
    print(f"Copia sincronizada: {REPO_WORD_PATH}")


if __name__ == "__main__":
    main()
