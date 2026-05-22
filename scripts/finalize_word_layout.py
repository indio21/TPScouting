from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
WORD_PATH = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_WORD_PATH = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"


def norm(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_text(paragraph: Paragraph) -> str:
    return norm(paragraph.text)


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro parrafo exacto: {exact}")


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._element is target._element:
            return index
    raise ValueError("No se encontro el parrafo en el documento.")


def remove_generated_section_breaks(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if not paragraph_text(paragraph) and paragraph._element.xpath("./w:pPr/w:sectPr"):
            paragraph._element.getparent().remove(paragraph._element)


def move_resumen_after_lists(doc: Document) -> None:
    resumen = find_paragraph(doc, "RESUMEN")
    toc = find_paragraph(doc, "ÍNDICE")
    lista_figuras = find_paragraph(doc, "LISTA DE FIGURAS")
    intro = find_paragraph(doc, "1. INTRODUCCIÓN")

    if paragraph_index(doc, resumen) > paragraph_index(doc, toc):
        return

    blocks = []
    node = resumen._element
    while node is not None and node is not toc._element:
        nxt = node.getnext()
        blocks.append(node)
        node.getparent().remove(node)
        node = nxt

    # Recalcular despues de mover los bloques anteriores.
    destination = find_paragraph(doc, "Tabla 9-3. Glosario técnico mínimo.")
    if paragraph_index(doc, destination) > paragraph_index(doc, intro):
        destination = lista_figuras

    anchor = destination._element
    for block in blocks:
        anchor.addnext(block)
        anchor = block


def reset_toc_area(doc: Document) -> None:
    toc = find_paragraph(doc, "ÍNDICE")
    lista_figuras = find_paragraph(doc, "LISTA DE FIGURAS")

    node = toc._element.getnext()
    while node is not None and node is not lista_figuras._element:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt

    placeholder = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "__TOC_PLACEHOLDER__"
    run.append(text)
    placeholder.append(run)
    toc._element.addnext(placeholder)


def base_section_properties(doc: Document) -> list[OxmlElement]:
    sect_pr = doc.sections[0]._sectPr
    allowed = {qn("w:pgSz"), qn("w:pgMar"), qn("w:cols"), qn("w:docGrid")}
    return [deepcopy(child) for child in sect_pr if child.tag in allowed]


def insert_section_break_before(doc: Document, paragraph: Paragraph, props: list[OxmlElement]) -> None:
    previous = paragraph._element.getprevious()
    if previous is not None and previous.tag == qn("w:p") and previous.xpath("./w:pPr/w:sectPr"):
        return

    break_paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sect_pr.append(type_el)
    for child in props:
        sect_pr.append(deepcopy(child))
    p_pr.append(sect_pr)
    break_paragraph.append(p_pr)
    paragraph._element.addprevious(break_paragraph)


def add_section_breaks(doc: Document) -> None:
    props = base_section_properties(doc)
    headings = [
        "ÍNDICE",
        "LISTA DE FIGURAS",
        "RESUMEN",
        "1. INTRODUCCIÓN",
        "2. MARCO TEÓRICO",
        "3. METODOLOGÍA",
        "4. ANÁLISIS Y DISEÑO",
        "5. DESARROLLO",
        "6. PRUEBAS Y RESULTADOS",
        "7. CONCLUSIÓN",
        "8. BIBLIOGRAFÍA",
        "9. ANEXOS TÉCNICOS DEL MVP",
        "10. ANEXOS VISUALES Y EVIDENCIA",
    ]
    for heading in headings:
        insert_section_break_before(doc, find_paragraph(doc, heading), props)


def update_bibliography_dates(doc: Document) -> None:
    replacements = {
        "PyTorch Foundation.": "• PyTorch Foundation. (s. f.). PyTorch documentation. Recuperado el 15 de mayo de 2026 de https://pytorch.org/docs/",
        "Pallets Projects.": "• Pallets Projects. (s. f.). Flask documentation. Recuperado el 16 de mayo de 2026 de https://flask.palletsprojects.com/",
        "SQLAlchemy.": "• SQLAlchemy. (s. f.). SQLAlchemy documentation. Recuperado el 17 de mayo de 2026 de https://docs.sqlalchemy.org/",
        "OWASP Foundation.": "• OWASP Foundation. (s. f.). OWASP Top 10:2021. Recuperado el 18 de mayo de 2026 de https://owasp.org/Top10/",
        "Hudl.": "• Hudl. (s. f.). Wyscout. Recuperado el 19 de mayo de 2026 de https://www.hudl.com/products/wyscout",
        "FIFA.": "• FIFA. (s. f.). Talent Development. Recuperado el 20 de mayo de 2026 de https://publications.fifa.com/es/talent-development/",
    }
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        for prefix, replacement in replacements.items():
            if text.startswith(f"• {prefix}"):
                paragraph.clear()
                paragraph.add_run(replacement)
                break


def main() -> None:
    doc = Document(WORD_PATH)
    remove_generated_section_breaks(doc)
    move_resumen_after_lists(doc)
    reset_toc_area(doc)
    update_bibliography_dates(doc)
    add_section_breaks(doc)
    doc.save(WORD_PATH)
    copy2(WORD_PATH, REPO_WORD_PATH)
    print(f"Layout base actualizado: {WORD_PATH}")
    print(f"Copia sincronizada: {REPO_WORD_PATH}")


if __name__ == "__main__":
    main()
