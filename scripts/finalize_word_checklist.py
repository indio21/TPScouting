from __future__ import annotations

import json
import math
import struct
import sys
import zlib
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
WORD_PATH = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_WORD_PATH = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"
METADATA_PATH = ROOT / "scouting_app" / "training_metadata.json"
EVIDENCE_DIR = ROOT / "docs" / "evidencia_word_render"
TRAINING_CURVE_PATH = EVIDENCE_DIR / "07_training_curve_real.png"


def norm(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_text(paragraph: Paragraph) -> str:
    return norm(paragraph.text)


def clear_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def block_after(cursor) -> OxmlElement:
    if isinstance(cursor, Paragraph):
        return cursor._element
    if isinstance(cursor, Table):
        return cursor._tbl
    return cursor


def paragraph_after(doc: Document, cursor, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_after(cursor).addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def paragraph_before(doc: Document, cursor, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_after(cursor).addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro parrafo exacto: {exact}")


def find_contains(doc: Document, fragment: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if fragment in paragraph_text(paragraph):
            return paragraph
    raise ValueError(f"No se encontro fragmento: {fragment}")


def table_after(doc: Document, cursor, rows: list[list[str]], headers: list[str]) -> Table:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    table_xml = table._tbl
    table_xml.getparent().remove(table_xml)
    block_after(cursor).addnext(table_xml)
    return table


def add_picture_before(doc: Document, cursor, image_path: Path, width: float = 6.0) -> Paragraph:
    paragraph = paragraph_before(doc, cursor, "")
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    return paragraph


def png_chunk(kind: bytes, data: bytes) -> bytes:
    return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)


def write_png(path: Path, width: int, height: int, pixels: bytearray) -> None:
    rows = bytearray()
    stride = width * 3
    for y in range(height):
        rows.append(0)
        start = y * stride
        rows.extend(pixels[start : start + stride])
    data = b"\x89PNG\r\n\x1a\n"
    data += png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    data += png_chunk(b"IDAT", zlib.compress(bytes(rows), 9))
    data += png_chunk(b"IEND", b"")
    path.write_bytes(data)


FONT = {
    "0": ["111", "101", "101", "101", "111"],
    "1": ["010", "110", "010", "010", "111"],
    "2": ["111", "001", "111", "100", "111"],
    "3": ["111", "001", "111", "001", "111"],
    "4": ["101", "101", "111", "001", "001"],
    "5": ["111", "100", "111", "001", "111"],
    "6": ["111", "100", "111", "101", "111"],
    "7": ["111", "001", "010", "010", "010"],
    "8": ["111", "101", "111", "101", "111"],
    "9": ["111", "101", "111", "001", "111"],
    ".": ["000", "000", "000", "000", "010"],
    "-": ["000", "000", "111", "000", "000"],
    "%": ["101", "001", "010", "100", "101"],
    "A": ["111", "101", "111", "101", "101"],
    "C": ["111", "100", "100", "100", "111"],
    "D": ["110", "101", "101", "101", "110"],
    "E": ["111", "100", "111", "100", "111"],
    "F": ["111", "100", "111", "100", "100"],
    "H": ["101", "101", "111", "101", "101"],
    "I": ["111", "010", "010", "010", "111"],
    "L": ["100", "100", "100", "100", "111"],
    "N": ["101", "111", "111", "111", "101"],
    "O": ["111", "101", "101", "101", "111"],
    "P": ["111", "101", "111", "100", "100"],
    "R": ["111", "101", "111", "110", "101"],
    "S": ["111", "100", "111", "001", "111"],
    "T": ["111", "010", "010", "010", "010"],
    "U": ["101", "101", "101", "101", "111"],
    "V": ["101", "101", "101", "101", "010"],
}


def generate_training_curve() -> None:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    history = metadata["pytorch"]["history"]
    width, height = 1200, 700
    pixels = bytearray([255] * width * height * 3)

    def set_pixel(x: int, y: int, color: tuple[int, int, int]) -> None:
        if 0 <= x < width and 0 <= y < height:
            index = (y * width + x) * 3
            pixels[index : index + 3] = bytes(color)

    def line(x0: int, y0: int, x1: int, y1: int, color: tuple[int, int, int], thickness: int = 2) -> None:
        dx = abs(x1 - x0)
        dy = -abs(y1 - y0)
        sx = 1 if x0 < x1 else -1
        sy = 1 if y0 < y1 else -1
        err = dx + dy
        x, y = x0, y0
        while True:
            for tx in range(-thickness, thickness + 1):
                for ty in range(-thickness, thickness + 1):
                    if tx * tx + ty * ty <= thickness * thickness:
                        set_pixel(x + tx, y + ty, color)
            if x == x1 and y == y1:
                break
            e2 = 2 * err
            if e2 >= dy:
                err += dy
                x += sx
            if e2 <= dx:
                err += dx
                y += sy

    def rect(x: int, y: int, w: int, h: int, color: tuple[int, int, int]) -> None:
        for yy in range(y, y + h):
            for xx in range(x, x + w):
                set_pixel(xx, yy, color)

    def text(x: int, y: int, value: str, color: tuple[int, int, int], scale: int = 3) -> None:
        cursor = x
        for char in value.upper():
            if char == " ":
                cursor += 4 * scale
                continue
            glyph = FONT.get(char)
            if not glyph:
                cursor += 4 * scale
                continue
            for gy, row in enumerate(glyph):
                for gx, bit in enumerate(row):
                    if bit == "1":
                        rect(cursor + gx * scale, y + gy * scale, scale, scale, color)
            cursor += 4 * scale

    left, top, right, bottom = 110, 80, 1120, 590
    axis = (35, 45, 55)
    grid = (225, 231, 235)
    colors = {
        "LOSS": (31, 119, 180),
        "VAL PR-AUC": (214, 39, 40),
        "VAL F1": (44, 160, 44),
    }
    y_min, y_max = 0.40, 0.70
    x_min, x_max = 1, max(item["epoch"] for item in history)

    text(330, 25, "CURVA REAL DE ENTRENAMIENTO", (20, 30, 40), 5)
    text(420, 55, "FUENTE TRAINING_METADATA.JSON", (80, 90, 100), 3)

    for tick in [0.40, 0.50, 0.60, 0.70]:
        y = bottom - int((tick - y_min) / (y_max - y_min) * (bottom - top))
        line(left, y, right, y, grid, 1)
        text(35, y - 8, f"{tick:.2f}", axis, 3)
    for tick in [1, 5, 10, 15]:
        x = left + int((tick - x_min) / (x_max - x_min) * (right - left))
        line(x, top, x, bottom, grid, 1)
        text(x - 12, bottom + 20, str(tick), axis, 3)

    line(left, bottom, right, bottom, axis, 2)
    line(left, top, left, bottom, axis, 2)
    text(560, 635, "EPOCH", axis, 4)

    series = [
        ("LOSS", [item["loss"] for item in history]),
        ("VAL PR-AUC", [item["val_pr_auc"] for item in history]),
        ("VAL F1", [item["val_f1"] for item in history]),
    ]
    epochs = [item["epoch"] for item in history]
    for name, values in series:
        points = []
        for epoch, value in zip(epochs, values):
            x = left + int((epoch - x_min) / (x_max - x_min) * (right - left))
            y = bottom - int((value - y_min) / (y_max - y_min) * (bottom - top))
            points.append((x, y))
        for (x0, y0), (x1, y1) in zip(points, points[1:]):
            line(x0, y0, x1, y1, colors[name], 3)
        for x, y in points:
            rect(x - 4, y - 4, 9, 9, colors[name])

    legend_x, legend_y = 775, 95
    for index, name in enumerate(["LOSS", "VAL PR-AUC", "VAL F1"]):
        y = legend_y + index * 34
        rect(legend_x, y, 28, 14, colors[name])
        text(legend_x + 40, y - 2, name, axis, 3)

    TRAINING_CURVE_PATH.parent.mkdir(parents=True, exist_ok=True)
    write_png(TRAINING_CURVE_PATH, width, height, pixels)


def insert_preliminaries(doc: Document) -> None:
    if any(paragraph_text(p) == "RESUMEN" for p in doc.paragraphs):
        return
    index_heading = find_paragraph(doc, "ÍNDICE")
    current = paragraph_before(doc, index_heading, "", None)
    current = paragraph_before(doc, current, "Palabras clave: scouting juvenil; fútbol; aprendizaje automático; Flask; PyTorch.", None)
    current = paragraph_before(doc, current, "Este Trabajo Final presenta TPScouting, un MVP académico para apoyar el seguimiento de jugadores juveniles de fútbol. La aplicación permite registrar jugadores, atributos, estadísticas, reportes y disponibilidad; visualizar información en un panel general; comparar perfiles; y calcular una proyección de potencial mediante un modelo implementado en PyTorch. El sistema utiliza Flask, SQLAlchemy, Jinja2, Bootstrap, Chart.js, pytest y despliegue en Render con Gunicorn. La validación se realiza con datos sintéticos, pruebas automatizadas, métricas de Machine Learning, comparación contra baseline y evidencia real del despliegue. El alcance se delimita como MVP académico y se documentan limitaciones: no reemplaza el criterio del scout, no integra video/tracking comercial y requiere datos reales longitudinales para validación productiva.", None)
    paragraph_before(doc, current, "RESUMEN", "Heading 1")


def remove_empty_headings(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.style.name.startswith("Heading") and not paragraph_text(paragraph):
            remove_paragraph(paragraph)


def insert_figure_and_table_lists(doc: Document) -> None:
    if any(paragraph_text(p) == "LISTA DE FIGURAS" for p in doc.paragraphs):
        return
    intro = find_paragraph(doc, "1. INTRODUCCIÓN")
    table_items = [
        "Tabla 3-1. Desglose de tareas y cronograma consolidado.",
        "Tabla 4-1. Comparación entre Flask y .NET Core Minimal APIs.",
        "Tabla 4-2. Comparación entre SQLAlchemy y Peewee.",
        "Tabla 4-3. Comparación entre motores relacionales.",
        "Tabla 4-4. Dimensiones de evaluación deportiva.",
        "Tabla 4-5. Capas y responsabilidades del MVP.",
        "Tabla 4-6. Requisitos funcionales del sistema.",
        "Tabla 4-7. Entidades principales del modelo de datos.",
        "Tabla 4-8. Endpoints principales del MVP.",
        "Tabla 4-9. Decisiones tecnológicas.",
        "Tabla 4-10. Casos de uso principales.",
        "Tabla 5-1. Configuración verificada del modelo PlayerNet.",
        "Tabla 5-2. Estado del despliegue en Render.",
        "Tabla 6-1. Cobertura del plan de pruebas.",
        "Tabla 6-2. Resultados de validación automatizada y CI.",
        "Tabla 6-3. Métricas de PlayerNet en test.",
        "Tabla 6-4. Comparación con baselines.",
        "Tabla 6-5. Smoke real del deploy en Render.",
        "Tabla 7-1. Cumplimiento de objetivos.",
        "Tabla 9-1. Rutas principales.",
        "Tabla 9-2. Variables de entorno relevantes.",
        "Tabla 9-3. Glosario técnico mínimo.",
    ]
    figure_items = [
        "Figura 6-1. Curva real de entrenamiento desde training_metadata.json.",
        "Figura 6-2. Login del MVP desplegado en Render.",
        "Figura 6-3. Panel general o mesa de scouting.",
        "Figura 6-4. Listado paginado de jugadores.",
        "Figura 6-5. Ficha de jugador con historial.",
        "Figura 6-6. Vista de predicción de potencial.",
        "Figura 6-7. Comparador múltiple de jugadores.",
    ]
    current = paragraph_before(doc, intro, "", None)
    for item in reversed(table_items):
        current = paragraph_before(doc, current, item, None)
    current = paragraph_before(doc, current, "LISTA DE TABLAS", "Heading 1")
    for item in reversed(figure_items):
        current = paragraph_before(doc, current, item, None)
    paragraph_before(doc, current, "LISTA DE FIGURAS", "Heading 1")


def replace_texts(doc: Document) -> None:
    replacements = {
        "Entonces, utilizamos SQLAlchemy porque ofrece:": "Se adopta SQLAlchemy porque ofrece:",
        "Entonces concluimos que:": "Conclusión técnica sobre el ORM adoptado:",
        "La arquitectura actual no es una MLP plana con Sigmoid final": "La arquitectura actual no es una MLP plana con activación final interna",
    }
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        for old, new in replacements.items():
            if old in text:
                clear_paragraph(paragraph, paragraph.text.replace(old, new))


def insert_training_curve(doc: Document) -> None:
    if any("Figura 6-1. Curva real de entrenamiento generada" in paragraph_text(p) for p in doc.paragraphs):
        return
    target = find_paragraph(doc, "6.2.3 Evidencia operativa del deploy")
    caption = paragraph_before(
        doc,
        target,
        "Figura 6-1. Curva real de entrenamiento generada desde training_metadata.json.",
        "Caption",
    )
    add_picture_before(doc, caption, TRAINING_CURVE_PATH, width=6.0)
    paragraph_before(
        doc,
        caption,
        "La siguiente curva se construye con la historia real de entrenamiento registrada por el pipeline: pérdida de entrenamiento, PR-AUC de validación y F1 de validación por época. No es un gráfico ilustrativo.",
        None,
    )

    renumbers = {
        "Figura 6-1. Login del MVP desplegado en Render.": "Figura 6-2. Login del MVP desplegado en Render.",
        "Figura 6-2. Panel general o mesa de scouting con métricas accionables.": "Figura 6-3. Panel general o mesa de scouting con métricas accionables.",
        "Figura 6-3. Listado paginado de jugadores con edad, categoría y potencial.": "Figura 6-4. Listado paginado de jugadores con edad, categoría y potencial.",
        "Figura 6-4. Ficha de jugador con historial, atributos y acciones CRUD en modales.": "Figura 6-5. Ficha de jugador con historial, atributos y acciones CRUD en modales.",
        "Figura 6-5. Vista de predicción de potencial del jugador.": "Figura 6-6. Vista de predicción de potencial del jugador.",
        "Figura 6-6. Comparador múltiple de jugadores.": "Figura 6-7. Comparador múltiple de jugadores.",
    }
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text in renumbers:
            clear_paragraph(paragraph, renumbers[text])


def insert_smoke_table(doc: Document) -> None:
    if any("Como evidencia operativa adicional, se ejecutó un smoke HTTP real" in paragraph_text(p) for p in doc.paragraphs):
        return
    target = find_contains(doc, "Las siguientes figuras fueron capturadas")
    intro = paragraph_before(
        doc,
        target,
        "Como evidencia operativa adicional, se ejecutó un smoke HTTP real contra Render el 20/05/2026. El primer request refleja arranque frío del plan Free; luego las rutas principales respondieron correctamente.",
        None,
    )
    smoke_table = table_after(
        doc,
        intro,
        [
            ["/health", "GET", "200", "109.36 s", "OK; arranque frío Render Free"],
            ["/login", "GET", "200", "0.39 s", "Pantalla de login renderizada"],
            ["/login", "POST", "200", "4.96 s", "Login admin OK; redirección final a /players"],
            ["/dashboard", "GET", "200", "7.86 s", "Panel general renderizado"],
            ["/players", "GET", "200", "0.38 s", "Listado de jugadores"],
            ["/compare", "GET", "200", "0.43 s", "Comparador 1v1"],
            ["/compare/multi", "GET", "200", "7.73 s", "Comparador múltiple"],
            ["/settings", "GET", "200", "0.44 s", "Configuración protegida"],
            ["/players/import", "GET", "200", "0.39 s", "Carga masiva"],
        ],
        ["Ruta", "Método", "Status", "Tiempo", "Resultado"],
    )
    return smoke_table


def add_development_detail(doc: Document) -> None:
    if any("Flujo técnico de inferencia" in paragraph_text(p) for p in doc.paragraphs):
        return
    backend_heading = find_paragraph(doc, "5.2 Implementación del frontend")
    current = paragraph_before(
        doc,
        backend_heading,
        "Flujo técnico de inferencia: el usuario ingresa a la vista de predicción de un jugador; Flask recupera el registro mediante SQLAlchemy; el preprocesador transforma atributos, edad, posición e historial en un vector de entrada; PlayerNet calcula logits; se aplica sigmoid para obtener una probabilidad; y la aplicación combina esa señal con reglas operativas e historial para mostrar una categoría comprensible al usuario.",
        None,
    )
    paragraph_before(
        doc,
        current,
        "La integración con PyTorch se mantiene dentro del backend Flask. Los artefactos model.pt, preprocessor.joblib y probability_calibrator.joblib se cargan al iniciar la aplicación y se validan contra input_dim para evitar incompatibilidades entre entrenamiento e inferencia.",
        None,
    )

    ml_heading = find_paragraph(doc, "5.5 Seguridad y robustez implementada")
    current = paragraph_before(
        doc,
        ml_heading,
        "El pipeline de entrenamiento registra metadata junto con el modelo: versión de checkpoint, input_dim, seed, partición de datos, métricas de validación/test, baseline logístico, baseline simple y artefactos generados. Esto permite reproducir la evidencia y detectar si una corrida futura cambia el comportamiento del modelo.",
        None,
    )
    paragraph_before(
        doc,
        current,
        "La selección del modelo no se presenta como superioridad automática de deep learning. En la corrida documentada, LogisticRegression(class_weight=\"balanced\") obtiene métricas levemente superiores en ROC-AUC, PR-AUC y F1. Por eso se conserva como baseline metodológico obligatorio y se documenta que PlayerNet queda integrado como modelo operativo del MVP.",
        None,
    )

    deploy_heading = find_paragraph(doc, "6. PRUEBAS Y RESULTADOS")
    paragraph_before(
        doc,
        deploy_heading,
        "En Render se evita SQLite como almacenamiento productivo debido al filesystem efímero. El despliegue utiliza PostgreSQL administrado, variables de entorno, Gunicorn con un worker y dos threads, y un seed demo idempotente para asegurar que la demo pública arranque con datos verificables.",
        None,
    )


def expand_glossary(doc: Document) -> None:
    glossary_table = None
    for table in doc.tables:
        header = " | ".join(norm(cell.text) for cell in table.rows[0].cells)
        if header == "Término | Definición":
            glossary_table = table
            break
    if glossary_table is None:
        raise ValueError("No se encontro tabla de glosario.")
    existing = {norm(row.cells[0].text) for row in glossary_table.rows[1:]}
    additions = [
        ("SCRUM", "Marco de trabajo ágil utilizado para organizar iteraciones, planificación y revisión del avance."),
        ("Sprint", "Iteración de trabajo con objetivo acotado dentro de SCRUM."),
        ("Dashboard / panel general", "Vista de control con indicadores, listas accionables y gráficos para apoyar decisiones."),
        ("SQLAlchemy", "ORM de Python utilizado para mapear entidades del dominio a tablas relacionales."),
        ("PostgreSQL", "Motor relacional administrado utilizado en el despliegue público de Render."),
        ("SQLite", "Motor relacional liviano utilizado para desarrollo local y pruebas del MVP."),
        ("Logits", "Salida cruda del modelo antes de convertirla a probabilidad mediante sigmoid."),
        ("Sigmoid", "Función matemática que transforma logits en valores entre 0 y 1."),
        ("Baseline", "Modelo o regla de referencia usado para comparar el desempeño de una propuesta."),
        ("Checkpoint", "Archivo que guarda estado del modelo y metadata necesaria para cargarlo de forma compatible."),
        ("input_dim", "Cantidad de features esperadas por el modelo al momento de inferir."),
        ("Smoke test", "Prueba rápida de disponibilidad y rutas críticas del sistema desplegado."),
        ("CI", "Integración continua; ejecución automatizada de pruebas en GitHub Actions."),
        ("pytest-cov", "Extensión de pytest usada para medir cobertura de pruebas."),
    ]
    for term, definition in additions:
        if term in existing:
            continue
        cells = glossary_table.add_row().cells
        cells[0].text = term
        cells[1].text = definition


def add_table_captions(doc: Document) -> None:
    if any(p.style.name == "Caption" and paragraph_text(p).startswith("Tabla 3-1.") for p in doc.paragraphs):
        return
    captions = [
        "Tabla 3-1. Desglose de tareas y cronograma consolidado.",
        "Tabla 4-1. Comparación entre Flask y .NET Core Minimal APIs.",
        "Tabla 4-2. Comparación entre SQLAlchemy y Peewee.",
        "Tabla 4-3. Comparación entre motores relacionales.",
        "Tabla 4-4. Dimensiones de evaluación deportiva.",
        "Tabla 4-5. Capas y responsabilidades del MVP.",
        "Tabla 4-6. Requisitos funcionales del sistema.",
        "Tabla 4-7. Entidades principales del modelo de datos.",
        "Tabla 4-8. Endpoints principales del MVP.",
        "Tabla 4-9. Decisiones tecnológicas.",
        "Tabla 4-10. Casos de uso principales.",
        "Tabla 5-1. Configuración verificada del modelo PlayerNet.",
        "Tabla 5-2. Estado del despliegue en Render.",
        "Tabla 6-1. Cobertura del plan de pruebas.",
        "Tabla 6-2. Resultados de validación automatizada y CI.",
        "Tabla 6-3. Métricas de PlayerNet en test.",
        "Tabla 6-4. Comparación con baselines.",
        "Tabla 6-5. Smoke real del deploy en Render.",
        "Tabla 7-1. Cumplimiento de objetivos.",
        "Tabla 9-1. Rutas principales.",
        "Tabla 9-2. Variables de entorno relevantes.",
        "Tabla 9-3. Glosario técnico mínimo.",
    ]
    tables = list(doc.tables)
    if len(tables) != len(captions):
        raise ValueError(f"Cantidad de tablas inesperada: {len(tables)} vs {len(captions)}")
    for table, caption in zip(tables, captions):
        paragraph_after(doc, table, caption, "Caption")


def remove_duplicate_consecutive_captions(doc: Document) -> None:
    previous = ""
    for paragraph in list(doc.paragraphs):
        text = paragraph_text(paragraph)
        if paragraph.style.name == "Caption" and text == previous:
            remove_paragraph(paragraph)
            continue
        previous = text if paragraph.style.name == "Caption" else ""


def main() -> int:
    generate_training_curve()
    doc = Document(WORD_PATH)
    insert_preliminaries(doc)
    remove_empty_headings(doc)
    insert_figure_and_table_lists(doc)
    replace_texts(doc)
    insert_training_curve(doc)
    insert_smoke_table(doc)
    add_development_detail(doc)
    expand_glossary(doc)
    add_table_captions(doc)
    remove_duplicate_consecutive_captions(doc)
    doc.save(WORD_PATH)
    copy2(WORD_PATH, REPO_WORD_PATH)
    print(f"Word actualizado: {WORD_PATH}")
    print(f"Copia repo: {REPO_WORD_PATH}")
    print(f"Curva real: {TRAINING_CURVE_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
