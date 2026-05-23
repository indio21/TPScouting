from __future__ import annotations

import sys
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
WORD_PATH = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_WORD_PATH = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagramas" / "export"

BODY_FIGURES = [
    (
        "06_casos_uso_general.png",
        "Figura 4-3. Diagrama de casos de uso general.",
        "La Figura 4-3 resume los casos de uso principales del sistema, agrupando los roles de administrador, scout, director y el sistema de Machine Learning.",
        5.9,
    ),
    (
        "07_casos_uso_gestion_jugadores.png",
        "Figura 4-4. Diagrama de casos de uso de gestión de jugadores.",
        "La Figura 4-4 detalla los casos de uso asociados a alta, edición, consulta, carga masiva e historiales de jugadores.",
        4.25,
    ),
    (
        "08_casos_uso_analisis_decision.png",
        "Figura 4-5. Diagrama de casos de uso de análisis y decisión scout.",
        "La Figura 4-5 muestra los casos de uso orientados a consulta, comparación, panel general y predicción de potencial.",
        6.25,
    ),
]

ANNEX_HEADING = "10.5 Casos de uso ampliados"
ANNEX_FIGURES = [
    (
        "06_casos_uso_general.png",
        "Figura 10-6. Diagrama de casos de uso general ampliado.",
        "La Figura 10-6 presenta en tamaño ampliado la vista general de casos de uso del sistema.",
        6.6,
    ),
    (
        "07_casos_uso_gestion_jugadores.png",
        "Figura 10-7. Diagrama de casos de uso de gestión de jugadores ampliado.",
        "La Figura 10-7 presenta en tamaño ampliado los casos de uso del módulo de jugadores.",
        4.65,
    ),
    (
        "08_casos_uso_analisis_decision.png",
        "Figura 10-8. Diagrama de casos de uso de análisis y decisión scout ampliado.",
        "La Figura 10-8 presenta en tamaño ampliado los casos de uso de análisis, comparación y predicción.",
        6.9,
    ),
]

FIGURE_LIST_ENTRIES = [
    "Figura 4-1. Diagrama de componentes de TPScouting.",
    "Figura 4-2. Diagrama de clases del modelo de dominio.",
    "Figura 4-3. Diagrama de casos de uso general.",
    "Figura 4-4. Diagrama de casos de uso de gestión de jugadores.",
    "Figura 4-5. Diagrama de casos de uso de análisis y decisión scout.",
    "Figura 5-1. Secuencia de prediccion de potencial de jugador.",
    "Figura 5-2. Secuencia de carga del panel general.",
    "Figura 5-3. Diagrama de despliegue en Render.",
    "Figura 6-1. Curva real de entrenamiento desde training_metadata.json.",
    "Figura 6-2. Login del MVP desplegado en Render.",
    "Figura 6-3. Panel general o mesa de scouting.",
    "Figura 6-4. Listado paginado de jugadores.",
    "Figura 6-5. Ficha de jugador con historial.",
    "Figura 6-6. Vista de predicción de potencial.",
    "Figura 6-7. Comparador múltiple de jugadores.",
    "Figura 10-1. Diagrama de componentes de TPScouting ampliado.",
    "Figura 10-2. Diagrama de clases del modelo de dominio ampliado.",
    "Figura 10-3. Diagrama de secuencia de predicción ampliado.",
    "Figura 10-4. Diagrama de secuencia del panel general ampliado.",
    "Figura 10-5. Diagrama de despliegue en Render ampliado.",
    "Figura 10-6. Diagrama de casos de uso general ampliado.",
    "Figura 10-7. Diagrama de casos de uso de gestión de jugadores ampliado.",
    "Figura 10-8. Diagrama de casos de uso de análisis y decisión scout ampliado.",
]


def norm(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_text(paragraph: Paragraph) -> str:
    return norm(paragraph.text)


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro parrafo exacto: {exact}")


def find_caption(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Caption" and paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"No se encontro caption exacto: {exact}")


def paragraph_after(doc: Document, cursor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    cursor._element.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def update_figure_list(doc: Document) -> None:
    list_heading = find_paragraph(doc, "LISTA DE FIGURAS")
    table_heading = find_paragraph(doc, "LISTA DE TABLAS")
    node = list_heading._element.getnext()
    while node is not None and node is not table_heading._element:
        next_node = node.getnext()
        if node.tag.endswith("}p"):
            text = norm(Paragraph(node, doc._body).text)
            if text.startswith("Figura "):
                node.getparent().remove(node)
        node = next_node
    for entry in FIGURE_LIST_ENTRIES:
        new_p = OxmlElement("w:p")
        table_heading._element.addprevious(new_p)
        paragraph = Paragraph(new_p, doc._body)
        paragraph.add_run(entry)


def add_figure_after(
    doc: Document,
    cursor: Paragraph,
    image_name: str,
    caption: str,
    intro: str,
    width: float,
) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Caption" and paragraph_text(paragraph) == caption:
            return paragraph
    current = paragraph_after(doc, cursor, intro)
    image_paragraph = paragraph_after(doc, current)
    image_paragraph.alignment = 1
    image_paragraph.add_run().add_picture(str(DIAGRAM_DIR / image_name), width=Inches(width))
    return paragraph_after(doc, image_paragraph, caption, "Caption")


def insert_body_figures(doc: Document) -> None:
    current = find_caption(doc, "Tabla 4-10. Casos de uso principales.")
    for image_name, caption, intro, width in BODY_FIGURES:
        current = add_figure_after(doc, current, image_name, caption, intro, width)


def append_annex_figures(doc: Document) -> None:
    if any(paragraph_text(paragraph) == ANNEX_HEADING for paragraph in doc.paragraphs):
        return
    doc.add_page_break()
    doc.add_paragraph(ANNEX_HEADING, style="Heading 2")
    doc.add_paragraph(
        "Esta sección incorpora los diagramas de casos de uso en páginas independientes para facilitar su lectura."
    )
    for index, (image_name, caption, intro, width) in enumerate(ANNEX_FIGURES):
        if index > 0:
            doc.add_page_break()
        doc.add_paragraph(intro)
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = 1
        image_paragraph.add_run().add_picture(str(DIAGRAM_DIR / image_name), width=Inches(width))
        doc.add_paragraph(caption, style="Caption")


def main() -> None:
    doc = Document(WORD_PATH)
    update_figure_list(doc)
    insert_body_figures(doc)
    append_annex_figures(doc)
    doc.save(WORD_PATH)
    copy2(WORD_PATH, REPO_WORD_PATH)
    print(f"Casos de uso insertados en: {WORD_PATH}")
    print(f"Copia sincronizada: {REPO_WORD_PATH}")


if __name__ == "__main__":
    main()
