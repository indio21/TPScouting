from __future__ import annotations

import sys
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
WORD_PATH = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_WORD_PATH = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagramas" / "export"

ANNEX_HEADING = "10.4 Diagramas ampliados"

ANNEX_FIGURES = [
    (
        "03_componentes.png",
        "Figura 10-1. Diagrama de componentes de TPScouting ampliado.",
        "La Figura 10-1 presenta en tamaño ampliado los componentes del cliente, la aplicación Flask, servicios, persistencia y módulo de Machine Learning.",
    ),
    (
        "05_clases.png",
        "Figura 10-2. Diagrama de clases del modelo de dominio ampliado.",
        "La Figura 10-2 presenta en tamaño ampliado las clases principales definidas en el modelo SQLAlchemy y sus relaciones.",
    ),
    (
        "01_secuencia_prediccion.png",
        "Figura 10-3. Diagrama de secuencia de predicción ampliado.",
        "La Figura 10-3 presenta en tamaño ampliado el flujo de consulta de potencial de un jugador.",
    ),
    (
        "02_secuencia_dashboard.png",
        "Figura 10-4. Diagrama de secuencia del panel general ampliado.",
        "La Figura 10-4 presenta en tamaño ampliado el flujo de carga del panel general, incluyendo cache y consultas a base de datos.",
    ),
    (
        "04_despliegue.png",
        "Figura 10-5. Diagrama de despliegue en Render ampliado.",
        "La Figura 10-5 presenta en tamaño ampliado la relación entre navegador, Render, Gunicorn, Flask, PostgreSQL y artefactos ML.",
    ),
]


def norm(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_after(doc: Document, cursor, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    cursor._element.addnext(new_p)
    paragraph = docx_paragraph(new_p, doc)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def docx_paragraph(element, doc: Document):
    from docx.text.paragraph import Paragraph

    return Paragraph(element, doc._body)


def find_paragraph(doc: Document, exact: str):
    for paragraph in doc.paragraphs:
        if norm(paragraph.text) == exact:
            return paragraph
    raise ValueError(f"No se encontro parrafo exacto: {exact}")


def add_page_break(paragraph) -> None:
    paragraph.add_run().add_break(7)


def update_figure_list(doc: Document) -> None:
    list_heading = find_paragraph(doc, "LISTA DE FIGURAS")
    table_heading = find_paragraph(doc, "LISTA DE TABLAS")

    existing_entries = {norm(p.text) for p in doc.paragraphs}
    for _, caption, _ in ANNEX_FIGURES:
        if caption in existing_entries:
            continue
        new_p = OxmlElement("w:p")
        table_heading._element.addprevious(new_p)
        paragraph = docx_paragraph(new_p, doc)
        paragraph.add_run(caption)


def append_annex(doc: Document) -> None:
    if any(norm(paragraph.text) == ANNEX_HEADING for paragraph in doc.paragraphs):
        return

    doc.add_page_break()
    heading = doc.add_paragraph(ANNEX_HEADING, style="Heading 2")
    doc.add_paragraph(
        "Esta sección incorpora los diagramas técnicos en páginas independientes para facilitar su lectura. "
        "Las mismas figuras se mantienen previamente en los capítulos correspondientes como referencia contextual."
    )

    current = heading
    for index, (image_name, caption, intro) in enumerate(ANNEX_FIGURES):
        if index > 0:
            doc.add_page_break()
        doc.add_paragraph(intro)
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = 1
        image_paragraph.add_run().add_picture(str(DIAGRAM_DIR / image_name), width=Inches(7.0))
        doc.add_paragraph(caption, style="Caption")


def main() -> None:
    doc = Document(WORD_PATH)
    update_figure_list(doc)
    append_annex(doc)
    doc.save(WORD_PATH)
    copy2(WORD_PATH, REPO_WORD_PATH)
    print(f"Anexo de diagramas ampliados insertado en: {WORD_PATH}")
    print(f"Copia sincronizada: {REPO_WORD_PATH}")


if __name__ == "__main__":
    main()
