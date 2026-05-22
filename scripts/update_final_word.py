import sys
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
SRC = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Lo que se envio.docx")
OUT = Path(r"C:\Users\Usuario\Desktop\TRABAJO_FINAL - Version corregida TPScouting.docx")
REPO_OUT = ROOT / "docs" / "TRABAJO_FINAL_corregido_TPScouting.docx"
IMG_DIR = ROOT / "docs" / "evidencia_word_render"


def norm(text: str) -> str:
    return " ".join((text or "").split())


def update_run_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_paragraph_text(doc: Document, old: str, new: str) -> None:
    needle = norm(old)
    for paragraph in doc.paragraphs:
        current = norm(paragraph.text)
        if current == needle or needle in current:
            update_run_text(paragraph, new)
            return
    raise ValueError(f"No se encontro texto para reemplazar: {old[:80]}")


def make_editor(doc: Document):
    def find_para_exact(text: str) -> Paragraph:
        needle = norm(text)
        for paragraph in doc.paragraphs:
            if norm(paragraph.text) == needle:
                return paragraph
        raise ValueError(f"No se encontro parrafo: {text}")

    def clear_between(start_text: str, end_text: str) -> Paragraph:
        start = find_para_exact(start_text)
        end = find_para_exact(end_text)
        node = start._p.getnext()
        while node is not None and node is not end._p:
            nxt = node.getnext()
            node.getparent().remove(node)
            node = nxt
        return start

    def clear_from(start_text: str) -> Paragraph:
        start = find_para_exact(start_text)
        node = start._p.getnext()
        while node is not None:
            nxt = node.getnext()
            if node.tag != qn("w:sectPr"):
                node.getparent().remove(node)
            node = nxt
        return start

    def add_after(cursor, text: str = "", style: str | None = None) -> Paragraph:
        base = cursor._element if hasattr(cursor, "_element") else cursor._tbl if hasattr(cursor, "_tbl") else cursor
        new_p = OxmlElement("w:p")
        base.addnext(new_p)
        paragraph = Paragraph(new_p, doc._body)
        if style:
            paragraph.style = style
        paragraph.add_run(text)
        return paragraph

    def add_bullets(cursor, items: list[str]) -> Paragraph:
        current = cursor
        for item in items:
            current = add_after(current, f"• {item}")
        return current

    def add_numbered(cursor, items: list[str]) -> Paragraph:
        current = cursor
        for index, item in enumerate(items, start=1):
            current = add_after(current, f"{index}. {item}")
        return current

    def heading_after(cursor, text: str, level: int = 2) -> Paragraph:
        return add_after(cursor, text, f"Heading {level}")

    def add_table_after(cursor, rows: list[list[str]], headers: list[str] | None = None):
        table = doc.add_table(rows=1 if headers else 0, cols=len(headers or rows[0]))
        table.style = "Table Grid"
        if headers:
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        tbl = table._tbl
        parent = tbl.getparent()
        parent.remove(tbl)
        base = cursor._element if hasattr(cursor, "_element") else cursor._tbl if hasattr(cursor, "_tbl") else cursor
        base.addnext(tbl)
        return table

    def add_picture_after(cursor, image_path: Path, caption: str, width: float = 6.0) -> Paragraph:
        paragraph = add_after(cursor, "")
        paragraph.add_run().add_picture(str(image_path), width=Inches(width))
        return add_after(paragraph, caption, "Caption")

    return {
        "find": find_para_exact,
        "clear_between": clear_between,
        "clear_from": clear_from,
        "add": add_after,
        "bullets": add_bullets,
        "numbered": add_numbered,
        "heading": heading_after,
        "table": add_table_after,
        "picture": add_picture_after,
    }


def apply_updates() -> None:
    copy2(SRC, OUT)
    doc = Document(OUT)
    ed = make_editor(doc)

    direct_replacements = {
        "SQL y motores relacionales (SQLite en el MVP; SQL Server como alternativa)": "SQL y motores relacionales (SQLite local; PostgreSQL en despliegue; SQL Server como alternativa futura)",
        "Base de datos: SQLite en el MVP, gestionada mediante SQLAlchemy (ORM). Se considera escalabilidad futura hacia motores como SQL Server/PostgreSQL sin afectar la lógica de dominio.": "Base de datos: SQLite para desarrollo local y PostgreSQL administrado en Render para la demo pública, ambos gestionados mediante SQLAlchemy (ORM). Se mantiene portabilidad hacia motores relacionales alternativos sin modificar la lógica de dominio principal.",
        "Almacenamiento: En el MVP, los datos históricos de los jugadores se almacenan en SQLite, estructurados para registrar atributos, evaluaciones y progresiones a lo largo del tiempo. En escenarios de crecimiento (más usuarios, mayor volumen o analítica avanzada), se contempla la migración a un motor como SQL Server o PostgreSQL.": "Almacenamiento: En desarrollo local los datos se pueden almacenar en SQLite; en el despliegue verificado se utiliza PostgreSQL en Render. La estructura registra jugadores, atributos, evaluaciones, disponibilidad, rendimiento y reportes a lo largo del tiempo.",
        "Para gestionar la persistencia de datos en el backend se utiliza SQLAlchemy ORM con integración directa en Flask, lo que abstrae el acceso al motor relacional y permite operar sobre SQLite en el MVP, manteniendo portabilidad hacia otros motores si se requiriera.": "Para gestionar la persistencia de datos en el backend se utiliza SQLAlchemy ORM con integración directa en Flask, lo que abstrae el acceso al motor relacional y permite operar con SQLite local y PostgreSQL en Render manteniendo portabilidad.",
        "Entonces, utilizamos SQLAlchemy porque ofrece:": "Se adopta SQLAlchemy porque ofrece:",
        "Entonces concluimos que:": "Conclusión técnica sobre el ORM adoptado:",
        "En el MVP se utiliza SQLite por su simplicidad de despliegue, bajo costo operativo y adecuación a un entorno académico/prototipo. La arquitectura se mantiene desacoplada a través de SQLAlchemy, lo que habilita una migración futura a SQL Server u otro motor relacional si el proyecto evoluciona hacia un escenario de mayor escala y concurrencia.": "En desarrollo local se utiliza SQLite por su simplicidad y bajo costo operativo. En el despliegue de Render se utiliza PostgreSQL administrado para evitar depender de archivos SQLite en un filesystem efímero. SQLAlchemy permite mantener el mismo modelo de dominio en ambos entornos.",
        "En términos de escalabilidad, seguridad y administración, motores como SQL Server o PostgreSQL pueden resultar apropiados en etapas posteriores. Sin embargo, este trabajo prioriza coherencia con el MVP y reproducibilidad del entorno, por lo que la base de datos se mantiene en SQLite durante la validación del alcance definido.": "Para una producción real con más usuarios, concurrencia o auditoría, se recomienda sostener PostgreSQL u otro motor administrado, junto con backups, monitoreo y migraciones formales. Para el MVP académico se documenta la diferencia entre entorno local y despliegue.",
    }
    for old, new in direct_replacements.items():
        replace_paragraph_text(doc, old, new)

    add = ed["add"]
    bullets = ed["bullets"]
    numbered = ed["numbered"]
    heading = ed["heading"]
    table = ed["table"]
    picture = ed["picture"]
    clear_between = ed["clear_between"]
    clear_from = ed["clear_from"]

    # 1.5 Alcance
    cur = clear_between("1.5 Alcance", "2. MARCO TEÓRICO")
    cur = add(cur, "El alcance del proyecto es el desarrollo de un MVP académico de scouting juvenil que permite registrar jugadores, cargar atributos y estadísticas, consultar fichas, comparar perfiles y calcular una proyección de potencial mediante un modelo de Machine Learning integrado al backend Flask.")
    cur = add(cur, "La solución se enfoca en una institución o entorno de evaluación acotado, con datos sintéticos o semisintéticos para demostrar el flujo completo sin depender de proveedores externos de datos deportivos.")
    cur = add(cur, "Incluye:")
    cur = bullets(cur, [
        "Autenticación por usuario y contraseña, roles básicos y protección de rutas sensibles.",
        "Gestión de jugadores, staff, atributos, rendimiento, disponibilidad, evaluaciones físicas, reportes y comparadores.",
        "Carga individual y carga masiva por CSV compatible con Excel, con previsualización y validación por fila.",
        "Dashboard o panel general con métricas accionables adaptadas al rol del usuario.",
        "Inferencia de potencial con PlayerNet en PyTorch, usando features derivadas de atributos, edad, posición e historial disponible.",
        "Pruebas automatizadas con pytest, cobertura en CI y smoke real sobre Render.",
    ])
    cur = add(cur, "Queda fuera del alcance del MVP:")
    cur = bullets(cur, [
        "Análisis automático de video, tracking GPS, integración con Wyscout/Hudl u otros proveedores comerciales.",
        "Gestión multi-club, módulos financieros, contratos, valuación económica o scouting internacional distribuido.",
        "Entrenamiento con datos reales longitudinales de clubes, ya que el MVP utiliza datos sintéticos para reproducibilidad académica.",
        "Arquitectura productiva multi-instancia con cache distribuido, rate limiting centralizado, auditoría completa y orquestación de pipelines.",
    ])
    cur = add(cur, "Criterios mínimos de éxito observables:")
    cur = bullets(cur, [
        "Ejecutar la suite automatizada sin fallos.",
        "Acceder al deploy público, iniciar sesión y navegar dashboard, listado, ficha, predicción y comparadores.",
        "Mantener atributos dentro de escala 1-20 y potencial categorizado en bajo, medio y alto con umbrales documentados.",
        "Registrar evidencia real de métricas de entrenamiento, smoke de Render y capturas de funcionamiento.",
    ])

    # 2.1.7 y trabajos relacionados
    cur = clear_between("2.1.7 Librerías de inteligencia artificial relevantes", "3. METODOLOGÍA")
    cur = heading(cur, "2.1.7.1 Redes neuronales feed-forward y clasificación binaria", 4)
    cur = add(cur, "Una red neuronal feed-forward transforma un vector de características en una salida numérica mediante capas densas y funciones de activación. En problemas de clasificación binaria, la salida puede interpretarse como un score que luego se transforma en probabilidad o categoría mediante un umbral definido.")
    cur = heading(cur, "2.1.7.2 Normalización y vector de características", 4)
    cur = add(cur, "La normalización permite llevar atributos heterogéneos a rangos comparables antes de ingresarlos al modelo. En scouting juvenil, los atributos pueden cargarse en una escala ordinal 1-20 y transformarse internamente a valores normalizados para el entrenamiento e inferencia.")
    cur = heading(cur, "2.1.7.3 Métricas de evaluación de clasificación", 4)
    cur = add(cur, "Para evaluar modelos de clasificación se emplean métricas complementarias. Accuracy mide aciertos globales, F1 combina precisión y recall, ROC-AUC evalúa separación entre clases y PR-AUC resulta especialmente útil cuando la clase positiva es minoritaria.")
    cur = heading(cur, "2.2 Trabajos relacionados y contexto de uso", 2)
    cur = add(cur, "Como referencia de dominio se consideran herramientas comerciales y marcos institucionales de detección de talento, como Wyscout/Hudl y publicaciones de FIFA sobre desarrollo de talentos. Estas soluciones muestran la relevancia del análisis de datos en fútbol, pero no sustituyen el objetivo académico del MVP: construir y explicar un flujo propio, reproducible y auditable.")
    cur = add(cur, "El presente trabajo no afirma competir con plataformas profesionales de video o big data deportivo. Su aporte se ubica en la integración controlada de una aplicación web, persistencia de datos, visualización, reglas de negocio y un modelo predictivo simple para apoyar la evaluación de jugadores juveniles.")

    # Capítulo 3
    cur = clear_between("3. METODOLOGÍA", "4. ANÁLISIS Y DISEÑO")
    cur = heading(cur, "3.1 Enfoque de datos y modelado", 2)
    cur = add(cur, "El MVP utiliza datos sintéticos y semisintéticos para validar el flujo completo de captura, persistencia, entrenamiento e inferencia. Esta decisión permite reproducibilidad y evita depender de bases privadas de clubes o proveedores externos durante la evaluación académica.")
    cur = heading(cur, "3.1.1 Recolección y estructura de datos", 3)
    cur = add(cur, "Los datos se organizan alrededor de jugadores juveniles, atributos técnicos, físicos, mentales, estadísticas de rendimiento, disponibilidad, reportes de scouting y evaluaciones físicas. La edad operativa y la categoría juvenil se derivan de la fecha de nacimiento; por ejemplo, un jugador nacido en 2010 pertenece a la categoría 2010.")
    cur = heading(cur, "3.1.2 Preprocesamiento y normalización", 3)
    cur = add(cur, "Los atributos deportivos se cargan en escala 1-20. Para el modelo se transforman a representaciones numéricas compatibles con el preprocesador. Las posiciones se codifican como features categóricas y se incorporan variables históricas cuando existen registros previos.")
    cur = heading(cur, "3.1.3 Entrenamiento y evaluación", 3)
    cur = add(cur, "El entrenamiento se realiza con semilla fija, partición train/validation/test y métricas de clasificación. El modelo operativo es PlayerNet en PyTorch. Scikit-learn se utiliza para partición de datos, métricas, baseline logístico y componentes de calibración, no como motor principal de inferencia productiva.")
    cur = heading(cur, "3.2 Metodología de desarrollo", 2)
    cur = add(cur, "Se aplicó un enfoque incremental compatible con SCRUM, organizando el desarrollo en bloques chicos y verificables. Cada cambio relevante se acompañó con pruebas focales, suite automatizada y documentación técnica para mantener trazabilidad.")
    cur = heading(cur, "3.3 Ciclo de vida del software", 2)
    cur = numbered(cur, [
        "Análisis de requisitos y alcance del MVP.",
        "Diseño de entidades, rutas, vistas y flujo de datos.",
        "Implementación iterativa en Flask, SQLAlchemy, Jinja2, Bootstrap, Chart.js y PyTorch.",
        "Pruebas automatizadas y validaciones manuales de flujos críticos.",
        "Despliegue en Render con Gunicorn y PostgreSQL administrado para demo académica.",
        "Documentación de limitaciones, evidencia y oportunidades de mejora.",
    ])
    cur = heading(cur, "3.4 Desglose de tareas y cronograma consolidado", 2)
    cur = table(cur, [
        ["1", "Análisis y requisitos", "Definición de alcance, actores, entidades y criterios de éxito.", "Completado"],
        ["2", "Backend y persistencia", "Modelos SQLAlchemy, rutas Flask, migraciones manuales y validaciones.", "Completado"],
        ["3", "Frontend y UX", "Vistas Jinja2, Bootstrap, modales CRUD, dashboard, comparadores y responsive.", "Completado"],
        ["4", "Machine Learning", "Dataset sintético, entrenamiento, métricas, checkpoint y carga en inferencia.", "Completado"],
        ["5", "Seguridad y robustez", "CSRF, sesión, roles, secret obligatoria, rate limiting básico y errores controlados.", "Completado para MVP"],
        ["6", "Testing y CI", "pytest, pytest-cov, GitHub Actions, smoke visual opcional y smoke Render.", "Completado"],
        ["7", "Documentación y entrega", "Alineación del trabajo escrito con la app real y evidencia verificable.", "En cierre"],
    ], ["Fase", "Bloque", "Actividad", "Estado"])

    # Capítulo 4 desde 4.2
    cur = clear_between("4.2 Funcionalidades y componentes", "5. DESARROLLO")
    cur = heading(cur, "4.2.1 Módulos funcionales implementados", 3)
    cur = bullets(cur, [
        "Autenticación, registro de usuarios y control básico por roles.",
        "Gestión de jugadores, ficha completa, edición, baja y carga masiva CSV.",
        "Historial de atributos, rendimiento, partidos, disponibilidad, evaluaciones físicas y reportes scout.",
        "Panel general dinámico por rol: mesa de scouting para scout/administrador y estado de plantel para director técnico.",
        "Comparador 1v1 y comparador múltiple con métricas visuales.",
        "Predicción de potencial y recomendaciones operativas en ficha de jugador.",
    ])
    cur = heading(cur, "4.2.2 Modelo de evaluación deportiva", 3)
    cur = add(cur, "La evaluación del jugador combina datos cuantitativos y cualitativos. Los atributos principales se registran en escala 1-20 y se agrupan en dimensiones técnicas, físicas, mentales y defensivas. El sistema no reemplaza el criterio del scout; aporta orden, trazabilidad y comparación consistente entre jugadores.")
    cur = table(cur, [
        ["Técnica", "pace, shooting, passing, dribbling, technique, vision", "Observación técnica y carga de atributos"],
        ["Física", "physical, stamina, strength, agility, pace", "Evaluaciones y reportes"],
        ["Defensiva", "defending, tackling, marking", "Perfil por posición"],
        ["Mental", "determination, decision_making, work_rate, composure", "Reportes y atributos"],
        ["Rendimiento", "minutos, goles, asistencias, precisión, puntaje final", "Historial de partidos/entrenamientos"],
    ], ["Dimensión", "Variables representativas", "Fuente"])
    cur = heading(cur, "4.2.3 Uso de PyTorch y Scikit-learn", 3)
    cur = add(cur, "PyTorch se emplea para implementar PlayerNet, el modelo que se carga en backend para inferencia. Scikit-learn se utiliza como soporte metodológico para partición de datos, métricas de evaluación, baseline de regresión logística balanceada y calibración. No se implementa una regresión lineal productiva para inferencia en la aplicación web.")
    cur = heading(cur, "4.2.4 Seguridad y autenticación", 3)
    cur = add(cur, "El MVP implementa autenticación por usuario y contraseña, contraseñas hasheadas con Werkzeug, sesiones firmadas por Flask, rutas protegidas, roles básicos y validación CSRF en formularios POST que modifican datos. En producción/Render la variable APP_SECRET_KEY es obligatoria y la aplicación no debe ejecutarse con claves por defecto.")
    cur = add(cur, "El login incorpora rate limiting en memoria por proceso. Esta mitigación es suficiente para una demo académica, pero no equivale a un limitador distribuido para producción multi-instancia.")
    cur = heading(cur, "4.2.5 Arquitectura lógica del MVP", 3)
    cur = add(cur, "La arquitectura actual es una aplicación Flask server-side con Jinja2 y Bootstrap en frontend, SQLAlchemy como ORM, SQLite para desarrollo local y PostgreSQL en Render para demo pública. El backend fue modularizado parcialmente mediante blueprints por familia: auth, staff, players, dashboard, compare y settings. La inferencia ML se encapsula en módulos de runtime, manteniendo compatibilidad con los endpoints existentes.")
    cur = table(cur, [
        ["Presentación", "templates/, static/", "Vistas Jinja2, Bootstrap, Chart.js, CSS propio"],
        ["Rutas", "scouting_app/routes/", "Blueprints por familia funcional"],
        ["Servicios", "scouting_app/services/", "Cache, seguridad, locks y datos operativos"],
        ["Dominio", "models.py, player_logic.py", "Entidades SQLAlchemy y reglas de negocio"],
        ["ML", "train_model.py, ml/runtime.py", "Entrenamiento, checkpoint, carga e inferencia"],
        ["Persistencia", "SQLite/PostgreSQL", "SQLite local; PostgreSQL en Render"],
    ], ["Capa", "Archivos principales", "Responsabilidad"])
    cur = heading(cur, "4.3 Requisitos del sistema", 2)
    cur = table(cur, [
        ["RF-01", "Autenticación por sesión para acceso a la aplicación.", "Implementado"],
        ["RF-02", "Gestión de usuarios con roles administrador, scout y director.", "Implementado"],
        ["RF-03", "Alta, edición y baja de jugadores.", "Implementado"],
        ["RF-04", "Ficha de jugador con perfil, edad, categoría, atributos, historial y predicción.", "Implementado"],
        ["RF-05", "Registro de atributos en escala 1-20 con historial.", "Implementado"],
        ["RF-06", "Registro de rendimiento, partidos, disponibilidad, físico y reportes.", "Implementado"],
        ["RF-07", "Cálculo de potencial con PlayerNet en PyTorch.", "Implementado"],
        ["RF-08", "Clasificación de potencial: bajo <60%, medio 60-79%, alto >=80%.", "Implementado"],
        ["RF-09", "Carga masiva CSV compatible con Excel y validación por fila.", "Implementado"],
        ["RF-10", "Panel general con métricas accionables por rol.", "Implementado"],
        ["RF-11", "Comparadores 1v1 y múltiple.", "Implementado"],
        ["RF-12", "Healthcheck y deploy en Render.", "Implementado"],
    ], ["ID", "Requisito", "Estado"])
    cur = heading(cur, "4.4 Modelo de datos", 2)
    cur = add(cur, "El modelo de datos se implementa con SQLAlchemy. La entidad central es Player, vinculada a historiales de atributos, estadísticas, partidos, disponibilidad, evaluaciones físicas y reportes scout. La fecha de nacimiento es la fuente para calcular edad actual y categoría juvenil.")
    cur = table(cur, [
        ["Player", "Jugador, datos personales, posición, club, atributos base, birth_date, photo_url, potencial."],
        ["PlayerAttributeHistory", "Historial de atributos técnicos/físicos/mentales en escala 1-20."],
        ["PlayerStat", "Estadísticas de rendimiento por fecha."],
        ["Match / PlayerMatchParticipation", "Contexto de partidos y participación del jugador."],
        ["PhysicalAssessment", "Evaluación física y mediciones asociadas."],
        ["PlayerAvailability", "Disponibilidad, fatiga y estado físico."],
        ["ScoutReport", "Reporte cualitativo del scout."],
        ["Coach / Director / User", "Staff, directivos y usuarios del sistema."],
    ], ["Entidad", "Propósito"])
    cur = heading(cur, "4.5 Endpoints principales del MVP", 2)
    cur = table(cur, [
        ["GET/POST", "/login", "Autenticación."],
        ["POST", "/logout", "Cierre de sesión con CSRF."],
        ["GET/POST", "/register", "Alta de usuarios según permisos."],
        ["GET", "/dashboard", "Panel general dinámico por rol."],
        ["GET", "/players", "Listado paginado y filtrable de jugadores."],
        ["GET/POST", "/players/manage", "Alta individual de jugador."],
        ["GET/POST", "/players/import", "Carga masiva CSV con previsualización."],
        ["GET/POST", "/player/<id>/stats", "Visualización y carga de rendimiento."],
        ["GET/POST", "/player/<id>/attributes", "Visualización y carga de atributos."],
        ["GET", "/player/<id>/predict", "Vista de proyección de potencial."],
        ["GET/POST", "/compare", "Comparador 1v1."],
        ["GET/POST", "/compare/multi", "Comparador múltiple."],
        ["GET/POST", "/settings", "Operaciones administrativas y pipeline."],
        ["GET", "/health", "Estado de servicio, base y calidad de datos."],
    ], ["Método", "Ruta", "Propósito"])
    cur = heading(cur, "4.6 Decisiones tecnológicas y alineación con el MVP", 2)
    cur = table(cur, [
        ["Flask", "Permite desarrollo web server-side simple, auditable y suficiente para MVP académico."],
        ["SQLAlchemy", "Abstrae SQLite local y PostgreSQL en Render sin reescribir reglas de dominio."],
        ["PyTorch", "Permite implementar PlayerNet y guardar checkpoints con metadata de entrada."],
        ["Scikit-learn", "Aporta métricas, partición, baseline y calibración."],
        ["Bootstrap/Jinja2", "Facilita UI funcional y responsive sin SPA compleja."],
        ["Chart.js", "Visualizaciones ligeras integradas en vistas HTML."],
        ["pytest/pytest-cov", "Validación automatizada y medición de cobertura en CI."],
        ["Render/Gunicorn", "Deploy público para smoke académico con variables reales y HTTPS."],
    ], ["Tecnología", "Justificación"])
    cur = heading(cur, "4.7 Casos de uso", 2)
    cur = table(cur, [
        ["CU-01", "Iniciar sesión", "Usuario registrado", "Sesión activa y acceso según rol."],
        ["CU-02", "Registrar jugador", "Scout o administrador", "Jugador persistido con fecha de nacimiento, edad y categoría."],
        ["CU-03", "Cargar atributos", "Jugador existente", "Historial actualizado en escala 1-20."],
        ["CU-04", "Cargar rendimiento/reportes", "Jugador existente", "Datos operativos visibles en ficha y dashboard."],
        ["CU-05", "Calcular potencial", "Modelo cargado y jugador con datos", "Probabilidad, categoría y sugerencias visibles."],
        ["CU-06", "Comparar jugadores", "Dos o más jugadores seleccionados", "Tabla/gráfico comparativo."],
        ["CU-07", "Revisar panel general", "Usuario autenticado", "KPIs y listas accionables por rol."],
    ], ["ID", "Caso de uso", "Precondición", "Resultado"])

    # Capítulo 5
    cur = clear_between("5. DESARROLLO", "6. PRUEBAS Y RESULTADOS")
    cur = heading(cur, "5.1 Implementación del backend", 2)
    cur = add(cur, "El backend se implementó en Flask con rutas server-side y blueprints por familia funcional. app.py conserva la inicialización, configuración, registro de blueprints, healthcheck, carga de artefactos ML y compatibilidad con endpoints históricos. Las rutas específicas se distribuyen en scouting_app/routes/auth.py, staff.py, players.py, dashboard.py, compare.py y settings.py.")
    cur = add(cur, "Se incorporaron servicios auxiliares para seguridad, cache, locks y datos operativos. Esta separación mejora la mantenibilidad sin convertir el MVP en una arquitectura de microservicios innecesaria para el alcance académico.")
    cur = heading(cur, "5.2 Implementación del frontend", 2)
    cur = add(cur, "El frontend utiliza Jinja2, Bootstrap, CSS propio y Chart.js. La UX fue ajustada para panel general, fichas de jugador, comparadores, formularios y modales CRUD. La carga de registros históricos se separó visualmente de la lectura de datos mediante modales/offcanvas, evitando formularios extensos incrustados en la pantalla principal.")
    cur = heading(cur, "5.3 Datos, edad y categoría juvenil", 2)
    cur = add(cur, "La fecha de nacimiento se consolidó como fuente de verdad. A partir de birth_date se calcula la edad actual y la categoría juvenil visible como “Cat. YYYY”. Para datos demo heredados sin fecha se generaron fechas determinísticas que conservan edades entre 12 y 18 años, dejando documentado que son datos sintéticos de MVP.")
    cur = heading(cur, "5.4 Entrenamiento del modelo de Machine Learning", 2)
    cur = add(cur, "El modelo operativo es PlayerNet. La arquitectura actual no es una MLP plana con Sigmoid final: combina una rama lineal amplia y una rama residual no lineal. El entrenamiento usa BCEWithLogitsLoss y la probabilidad se obtiene aplicando sigmoid sobre logits durante evaluación e inferencia.")
    cur = table(cur, [
        ["Clase", "PlayerNet"],
        ["input_dim", "68"],
        ["Loss", "BCEWithLogitsLoss"],
        ["Optimizador", "AdamW"],
        ["Learning rate", "0.0005"],
        ["Batch size", "256"],
        ["Épocas solicitadas / entrenadas", "45 / 15"],
        ["Seed", "42"],
        ["Checkpoint", "Incluye model_state, input_dim y versión"],
    ], ["Elemento", "Valor verificado"])
    cur = heading(cur, "5.5 Seguridad y robustez implementada", 2)
    cur = bullets(cur, [
        "APP_SECRET_KEY obligatoria en producción/Render.",
        "Protección CSRF en formularios POST mutantes.",
        "Logout por POST con CSRF.",
        "Passwords hasheadas con Werkzeug y validación mínima para usuarios nuevos.",
        "Rate limiting básico de login en memoria.",
        "Bloqueo de SQLite en producción salvo opt-in explícito.",
        "Healthcheck con validación de base y calidad mínima de datos.",
    ])
    cur = heading(cur, "5.6 Despliegue del sistema", 2)
    cur = add(cur, "El MVP se desplegó en Render con Gunicorn y PostgreSQL administrado. Por limitación del plan Free se utiliza una sola base PostgreSQL para la demo. La aplicación se valida mediante /health, login, panel general, listado, comparadores y vistas de jugador.")
    cur = table(cur, [
        ["URL pública", "https://tpscouting-mvp.onrender.com"],
        ["Servidor WSGI", "Gunicorn --workers 1 --threads 2"],
        ["Base en deploy", "PostgreSQL administrado Render Free"],
        ["Datos demo", "100 jugadores con birth_date, national_id y photo_url completos"],
        ["Modelo", "Artefactos versionados: model.pt, preprocessor.joblib, probability_calibrator.joblib"],
    ], ["Elemento", "Estado"])

    # Capítulo 6
    cur = clear_between("6. PRUEBAS Y RESULTADOS", "7. CONCLUSIÓN")
    cur = heading(cur, "6.1 Plan de pruebas", 2)
    cur = add(cur, "El plan de pruebas cubre flujos funcionales, seguridad, consistencia de datos, ML, dashboard, comparadores, carga masiva, CRUD de historiales y comportamiento sin modelo cargado. La suite automatizada se ejecuta con pytest y el CI genera cobertura mediante pytest-cov.")
    cur = table(cur, [
        ["Funcional", "Login, logout, rutas protegidas, jugadores, staff, comparadores, dashboard.", "pytest"],
        ["Seguridad", "CSRF, roles, password mínima, secret obligatoria, rate limiting login.", "pytest"],
        ["Datos", "birth_date, categoría, escala 1-20, migraciones manuales y validaciones.", "pytest"],
        ["ML", "checkpoint, input_dim, métricas, warnings y compatibilidad legacy.", "pytest"],
        ["UI/Smoke", "Render real y Playwright opcional.", "Smoke HTTP / capturas"],
    ], ["Tipo", "Cobertura", "Evidencia"])
    cur = heading(cur, "6.2 Resultados del MVP", 2)
    cur = add(cur, "La validación local más reciente ejecutada sobre el repositorio actual arrojó: 83 tests aprobados, 1 test omitido y 4 warnings conocidos de scikit-learn por valores NaN controlados en pruebas de preprocesamiento.")
    cur = table(cur, [
        ["pytest -q", "83 passed, 1 skipped, 4 warnings"],
        ["CI", "GitHub Actions ejecuta pytest con pytest-cov y sube coverage.xml como artefacto"],
        ["Smoke Render", "/login, /dashboard, /players, /compare, /compare/multi, /settings y /players/import respondieron correctamente en validación autenticada"],
        ["Healthcheck Render", "players_total=100; faltantes críticos en 0 según validación previa"],
    ], ["Evidencia", "Resultado"])
    cur = heading(cur, "6.2.1 Métricas del modelo PlayerNet", 3)
    cur = add(cur, "La corrida registrada en training_metadata.json corresponde al 19/05/2026, con semilla 42, dataset sintético de 20.000 jugadores y partición train/validation/test de 14.000/3.000/3.000.")
    cur = table(cur, [
        ["Accuracy", "0.9303"],
        ["ROC-AUC", "0.9174"],
        ["PR-AUC", "0.5241"],
        ["F1", "0.5282"],
        ["Precision", "0.5764"],
        ["Recall", "0.4875"],
        ["Matriz de confusión", "[[2674, 86], [123, 117]]"],
    ], ["Métrica PyTorch test", "Valor"])
    cur = heading(cur, "6.2.2 Comparación con baseline", 3)
    cur = add(cur, "Se conserva una comparación honesta con LogisticRegression(class_weight=\"balanced\") y un baseline simple por promedio de atributos. En esta corrida, el baseline logístico balanceado queda levemente por encima de PlayerNet en ROC-AUC, PR-AUC y F1; por lo tanto, el trabajo no afirma superioridad general del modelo PyTorch, sino integración funcional y trazabilidad del pipeline completo.")
    cur = table(cur, [
        ["PlayerNet PyTorch", "0.9303", "0.9174", "0.5241", "0.5282"],
        ["LogisticRegression balanced", "0.9310", "0.9205", "0.5378", "0.5327"],
        ["Promedio simple de atributos", "0.8960", "0.8390", "0.3513", "0.4201"],
    ], ["Modelo", "Accuracy", "ROC-AUC", "PR-AUC", "F1"])
    cur = heading(cur, "6.2.3 Evidencia operativa del deploy", 3)
    cur = add(cur, "Las siguientes figuras fueron capturadas desde el deploy público en Render, con sesión autenticada, para reemplazar los placeholders de la versión previa del documento.")
    for fname, caption in [
        ("01_login_render.png", "Figura 6-1. Login del MVP desplegado en Render."),
        ("02_dashboard_render.png", "Figura 6-2. Panel general o mesa de scouting con métricas accionables."),
        ("03_players_render.png", "Figura 6-3. Listado paginado de jugadores con edad, categoría y potencial."),
        ("04_player_detail_render.png", "Figura 6-4. Ficha de jugador con historial, atributos y acciones CRUD en modales."),
        ("05_prediction_render.png", "Figura 6-5. Vista de predicción de potencial del jugador."),
        ("06_compare_multi_render.png", "Figura 6-6. Comparador múltiple de jugadores."),
    ]:
        img = IMG_DIR / fname
        if img.exists():
            cur = picture(cur, img, caption, width=6.0)
    cur = heading(cur, "6.3 Limitaciones observadas en la validación", 2)
    cur = bullets(cur, [
        "Los datos son sintéticos; las métricas no representan desempeño real en un club.",
        "Render Free puede presentar arranque frío y demoras iniciales, especialmente en dashboard antes de cache.",
        "El cache, rate limiting y lock del pipeline son locales/in-memory; para producción real deberían externalizarse.",
        "El seed demo actual muestra pocos jugadores de alto potencial con el umbral >=80%; esto es coherente con la regla vigente y no debe maquillarse.",
    ])

    # Capítulo 7
    cur = clear_between("7. CONCLUSIÓN", "8. BIBLIOGRAFÍA")
    cur = heading(cur, "7.1 Conclusiones", 2)
    cur = add(cur, "El desarrollo del MVP demuestra la factibilidad de integrar una aplicación web de scouting juvenil con persistencia de datos, visualización, gestión de historiales y un modelo de Machine Learning para apoyar la evaluación de potencial. La solución es adecuada para demo académica y MVP, no para producción deportiva real sin nuevas etapas de validación, datos reales y endurecimiento operativo.")
    cur = add(cur, "El principal valor del proyecto no está en reemplazar la evaluación humana, sino en ordenar datos, hacer trazables los criterios y ofrecer una base técnica reproducible para comparar jugadores juveniles.")
    cur = heading(cur, "7.2 Cumplimiento de objetivos", 2)
    cur = table(cur, [
        ["Registrar jugadores y datos deportivos", "Cumplido", "CRUD de jugadores, atributos, rendimiento, reportes y carga masiva CSV."],
        ["Visualizar información para scouting", "Cumplido", "Dashboard, ficha, gráficos y comparadores."],
        ["Calcular potencial con ML", "Cumplido para MVP", "PlayerNet integrado, checkpoint con input_dim y métricas documentadas."],
        ["Asegurar acceso controlado", "Cumplido para MVP", "Login, roles, CSRF y secret obligatoria en producción."],
        ["Desplegar una demo pública", "Cumplido", "Render con PostgreSQL y smoke real."],
        ["Validar con datos reales deportivos", "Pendiente/futuro", "El MVP usa dataset sintético por reproducibilidad académica."],
    ], ["Objetivo", "Estado", "Evidencia"])
    cur = heading(cur, "7.3 Limitaciones", 2)
    cur = bullets(cur, [
        "Dataset sintético y ausencia de validación longitudinal con jugadores reales.",
        "Sin integración de video, tracking, GPS ni proveedores comerciales de datos.",
        "Persistencia y operación adecuadas para demo, pero no para producción multiusuario intensiva.",
        "Cache, rate limiting y locks no distribuidos.",
        "El modelo PyTorch no supera en todos los indicadores al baseline logístico en la corrida actual; se documenta como hallazgo metodológico.",
    ])
    cur = heading(cur, "7.4 Trabajo futuro", 2)
    cur = bullets(cur, [
        "Validar el modelo con datos reales y outcomes observables: minutos competitivos, promociones, consenso scout o firma profesional.",
        "Incorporar auditoría de cambios, logs estructurados y cache/rate limiting distribuido.",
        "Separar más lógica de app.py hacia servicios o una app factory si el proyecto crece.",
        "Agregar exportaciones formales, reportes PDF y flujos de aprobación por rol.",
        "Evaluar modelos alternativos y seleccionar el más adecuado según evidencia, no por complejidad técnica.",
    ])

    # Bibliografía
    cur = clear_between("8. BIBLIOGRAFÍA", "9. REFERENCIAS")
    cur = add(cur, "Las referencias externas se consolidan en esta sección. Los documentos internos del proyecto se trasladan a anexos para evitar confundir bibliografía con evidencia propia.")
    for item in [
        "Sommerville, I. (2011). Ingeniería de software (9.ª ed.). Pearson Educación.",
        "Pressman, R. S. (2010). Ingeniería del software: un enfoque práctico (7.ª ed.). McGraw-Hill.",
        "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
        "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning. Springer.",
        "PyTorch Foundation. (s. f.). PyTorch documentation. Recuperado el 19 de mayo de 2026 de https://pytorch.org/docs/",
        "Pallets Projects. (s. f.). Flask documentation. Recuperado el 19 de mayo de 2026 de https://flask.palletsprojects.com/",
        "SQLAlchemy. (s. f.). SQLAlchemy documentation. Recuperado el 19 de mayo de 2026 de https://docs.sqlalchemy.org/",
        "OWASP Foundation. (s. f.). OWASP Top 10:2021. Recuperado el 19 de mayo de 2026 de https://owasp.org/Top10/",
        "Hudl. (s. f.). Wyscout. Recuperado el 19 de mayo de 2026 de https://www.hudl.com/products/wyscout",
        "FIFA. (s. f.). Talent Development. Recuperado el 19 de mayo de 2026 de https://publications.fifa.com/es/talent-development/",
    ]:
        cur = add(cur, f"• {item}")

    # Anexos finales: se reemplaza Referencias internas + Anexos vacíos.
    chapter9 = ed["find"]("9. REFERENCIAS")
    update_run_text(chapter9, "9. ANEXOS TÉCNICOS DEL MVP")
    cur = clear_from("9. ANEXOS TÉCNICOS DEL MVP")
    cur = add(cur, "Esta sección reúne documentación propia del proyecto. No se presenta como bibliografía externa, sino como anexos técnicos verificables contra el código fuente actual.")
    cur = heading(cur, "9.1 Rutas principales", 2)
    cur = table(cur, [
        ["auth.py", "/login, /logout, /register"],
        ["players.py", "/players, /player/<id>, /players/manage, /players/import, historiales y predicción"],
        ["dashboard.py", "/dashboard"],
        ["compare.py", "/compare, /compare/multi"],
        ["staff.py", "/coaches, /directors"],
        ["settings.py", "/settings"],
        ["app.py", "/, /health y registro de blueprints"],
    ], ["Archivo", "Rutas relevantes"])
    cur = heading(cur, "9.2 Variables de entorno relevantes", 2)
    cur = table(cur, [
        ["APP_SECRET_KEY", "Obligatoria en producción para firmar sesiones."],
        ["APP_DB_URL", "Base operativa SQLite local o PostgreSQL en Render."],
        ["TRAINING_DB_URL", "Base de entrenamiento cuando se ejecuta pipeline."],
        ["ADMIN_USERNAME / ADMIN_PASSWORD", "Credenciales iniciales del administrador."],
        ["ALLOW_DEFAULT_ADMIN", "Solo desarrollo local; no usar en producción."],
        ["CACHE_TTL_SECONDS / CACHE_MAX_ENTRIES", "Cache in-memory configurable."],
        ["POTENTIAL_MEDIUM_THRESHOLD / POTENTIAL_HIGH_THRESHOLD", "Umbrales documentados; default 0.60 y 0.80."],
        ["RENDER_SMOKE_BASE_URL", "URL pública usada para smoke externo."],
    ], ["Variable", "Uso"])
    cur = heading(cur, "9.3 Glosario mínimo", 2)
    cur = table(cur, [
        ["MVP", "Producto mínimo viable utilizado para validar el flujo principal del sistema."],
        ["PlayerNet", "Modelo PyTorch integrado para estimar potencial de jugadores."],
        ["Feature vector", "Vector numérico de entrada al modelo."],
        ["ROC-AUC", "Métrica de separación entre clases para clasificación binaria."],
        ["PR-AUC", "Métrica útil con clase positiva minoritaria."],
        ["BCEWithLogitsLoss", "Función de pérdida usada con logits en clasificación binaria."],
        ["CSRF", "Ataque de falsificación de petición; mitigado con token en formularios POST."],
        ["Gunicorn", "Servidor WSGI usado para ejecutar Flask en Render."],
        ["Render", "Proveedor cloud utilizado para la demo pública."],
    ], ["Término", "Definición"])
    cur = heading(cur, "10. ANEXOS VISUALES Y EVIDENCIA", 1)
    cur = add(cur, "Los anexos visuales complementan la sección de pruebas con capturas reales del sistema desplegado y referencias a documentos internos del repositorio.")
    cur = heading(cur, "10.1 Evidencia Render", 2)
    cur = bullets(cur, [
        "URL pública validada: https://tpscouting-mvp.onrender.com.",
        "Login admin verificado en smoke real.",
        "Rutas verificadas: /dashboard, /players, /compare, /compare/multi, /settings y /players/import.",
        "Healthcheck previo documentado con 100 jugadores demo y faltantes críticos en 0.",
    ])
    cur = heading(cur, "10.2 Documentos técnicos internos del repositorio", 2)
    cur = bullets(cur, [
        "README.md: instalación, pruebas, deploy y limitaciones.",
        "RUNBOOK.md: operación, variables, smoke, Render y riesgos conocidos.",
        "docs/model_training_evidence.md: evidencia del entrenamiento y comparación de modelos.",
        "docs/auditoria_pendientes_2026-05-17.md: riesgos restantes y decisiones MVP.",
        "docs/cierre_pre_entrega_word_render_2026-05-18.md: cierre pre-entrega, deploy y smoke.",
    ])
    cur = heading(cur, "10.3 Limitaciones documentadas del MVP", 2)
    bullets(cur, [
        "Datos sintéticos y métricas no extrapolables a rendimiento real sin validación externa.",
        "Render Free puede dormir el servicio y producir latencias iniciales elevadas.",
        "Cache, rate limiting y lock de pipeline son locales al proceso.",
        "La app está modularizada por blueprints, pero todavía no es una arquitectura de microservicios.",
        "La decisión de modelo debe revisarse con datos reales; el baseline logístico fue competitivo en la corrida actual.",
    ])

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUT)
    copy2(OUT, REPO_OUT)
    print(f"saved {OUT} {OUT.stat().st_size}")
    print(f"saved {REPO_OUT} {REPO_OUT.stat().st_size}")


if __name__ == "__main__":
    apply_updates()
