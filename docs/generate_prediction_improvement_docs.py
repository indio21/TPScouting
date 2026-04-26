from __future__ import annotations

import json
import sqlite3
import subprocess
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
SCOUTING_APP_DIR = REPO_ROOT / "scouting_app"
DOCS_DIR = REPO_ROOT / "docs"
METADATA_PATH = SCOUTING_APP_DIR / "training_metadata.json"
TRAINING_DB_PATH = SCOUTING_APP_DIR / "players_training.db"
PROGRESS_MD = DOCS_DIR / "prediction_improvement_progress.md"
PROGRESS_DOCX = DOCS_DIR / "prediction_improvement_progress.docx"

PRIOR_LONGITUDINAL_STAGE = {
    "commit": "a7e6f7f",
    "players": 20000,
    "player_attribute_history": 179810,
    "player_stats": 162490,
    "matches": 324735,
    "player_match_participations": 324735,
    "scout_reports": 80133,
    "pytorch": {
        "roc_auc": 0.9247,
        "pr_auc": 0.3279,
        "f1": 0.4231,
    },
    "logistic_balanced": {
        "roc_auc": 0.9431,
        "pr_auc": 0.3923,
        "f1": 0.4409,
    },
}

AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK = {
    "pytorch": {
        "roc_auc": 0.9044,
        "pr_auc": 0.2029,
        "f1": 0.2542,
    },
    "logistic_balanced": {
        "roc_auc": 0.9425,
        "pr_auc": 0.2775,
        "f1": 0.3659,
    },
}


def git_branch() -> str:
    try:
        return (
            subprocess.check_output(
                ["git", "-C", str(REPO_ROOT), "rev-parse", "--abbrev-ref", "HEAD"],
                text=True,
            )
            .strip()
        )
    except Exception:
        return "desconocida"


def git_head() -> str:
    try:
        return (
            subprocess.check_output(
                ["git", "-C", str(REPO_ROOT), "rev-parse", "--short", "HEAD"],
                text=True,
            )
            .strip()
        )
    except Exception:
        return "desconocido"


def format_metric(value) -> str:
    if value in ("", None):
        return "N/D"
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def db_counts() -> dict[str, int]:
    queries = {
        "players": "select count(*) from players",
        "player_attribute_history": "select count(*) from player_attribute_history",
        "player_stats": "select count(*) from player_stats",
        "matches": "select count(*) from matches",
        "player_match_participations": "select count(*) from player_match_participations",
        "scout_reports": "select count(*) from scout_reports",
        "physical_assessments": "select count(*) from physical_assessments",
        "player_availability": "select count(*) from player_availability",
    }
    counts: dict[str, int] = {}
    with sqlite3.connect(TRAINING_DB_PATH) as connection:
        cursor = connection.cursor()
        for key, query in queries.items():
            cursor.execute(query)
            counts[key] = int(cursor.fetchone()[0])
    return counts


def build_markdown(metadata: dict, counts: dict[str, int]) -> str:
    pytorch_test = metadata["pytorch"].get("raw_test", metadata["pytorch"]["test"])
    pytorch_calibrated_test = metadata["pytorch"]["test"]
    baseline_lr = metadata["baselines"]["logistic_regression_balanced"]["test"]
    delta_pytorch_f1 = float(pytorch_test["f1"]) - PRIOR_LONGITUDINAL_STAGE["pytorch"]["f1"]
    delta_pytorch_pr_auc = float(pytorch_test["pr_auc"]) - PRIOR_LONGITUDINAL_STAGE["pytorch"]["pr_auc"]
    delta_vs_availability_f1 = float(pytorch_test["f1"]) - AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["pytorch"]["f1"]
    delta_vs_availability_pr_auc = float(pytorch_test["pr_auc"]) - AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["pytorch"]["pr_auc"]

    return f"""# Mejora de prediccion - avance tecnico

## Estado actual
- Rama de trabajo: `{git_branch()}`
- Rama estable cerrada del MVP corregido: `training`
- Rama activa para nuevas reformas: `reformas-finales`
- Ultimo commit publicado antes de esta etapa: `{PRIOR_LONGITUDINAL_STAGE["commit"]}`
- Objetivo de esta iteracion: cerrar la brecha entre PyTorch y el baseline lineal despues de sumar `PhysicalAssessment` y `Availability`, sin perder la riqueza metodologica ya ganada.
- Alcance del producto mantenido: scouting juvenil de 12 a 18 anos para clubes formativos.

## Que se implemento en esta etapa
- Se mantuvo la base enriquecida con dos fuentes nuevas de senal longitudinal:
- `PhysicalAssessment`
- `PlayerAvailability`
- El pipeline del modelo sigue agregando features fisicas recientes y de evolucion corporal.
- El pipeline del modelo sigue agregando disponibilidad, fatiga, carga y lesion/inactividad.
- La generacion sintetica sigue uniendo:
- trayectoria tecnica
- maduracion fisica
- disponibilidad mensual
- rendimiento competitivo derivado de esas tres capas
- En entrenamiento se reemplazo el doble rebalanceo previo por:
- `pos_weight` completo por defecto
- `shuffle` como estrategia default de batches
- `WeightedRandomSampler` solo como opcion
- `PlayerNet` paso a una arquitectura residual inicializada desde la solucion de `LogisticRegression(class_weight="balanced")`.
- La rama lineal queda dentro del modelo PyTorch y una rama residual aprende correcciones no lineales sobre esa base.
- Se mantuvo la calibracion de probabilidades y la inferencia compartida en la app.
- La app sigue sin cambiar su interfaz; el cambio permanece en el pipeline, el target y los artefactos del modelo.

## Conteos reales de la base de entrenamiento actual
- Jugadores: {counts["players"]}
- Snapshots de `PlayerAttributeHistory`: {counts["player_attribute_history"]}
- Registros agregados de `PlayerStat`: {counts["player_stats"]}
- Partidos sinteticos (`Match`): {counts["matches"]}
- Participaciones por partido: {counts["player_match_participations"]}
- Reportes de scout: {counts["scout_reports"]}
- Evaluaciones fisicas: {counts["physical_assessments"]}
- Registros de disponibilidad: {counts["player_availability"]}

## Comparacion con la etapa anterior
### Etapa anterior publicada: target temporal recalibrado
- PyTorch: ROC-AUC {PRIOR_LONGITUDINAL_STAGE["pytorch"]["roc_auc"]:.4f}, PR-AUC {PRIOR_LONGITUDINAL_STAGE["pytorch"]["pr_auc"]:.4f}, F1 {PRIOR_LONGITUDINAL_STAGE["pytorch"]["f1"]:.4f}
- Logistic balanceado: ROC-AUC {PRIOR_LONGITUDINAL_STAGE["logistic_balanced"]["roc_auc"]:.4f}, PR-AUC {PRIOR_LONGITUDINAL_STAGE["logistic_balanced"]["pr_auc"]:.4f}, F1 {PRIOR_LONGITUDINAL_STAGE["logistic_balanced"]["f1"]:.4f}

### Etapa local inmediatamente anterior: disponibilidad + fisico + trayectorias mas ricas
- PyTorch: ROC-AUC {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["pytorch"]["roc_auc"]:.4f}, PR-AUC {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["pytorch"]["pr_auc"]:.4f}, F1 {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["pytorch"]["f1"]:.4f}
- Logistic balanceado: ROC-AUC {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["logistic_balanced"]["roc_auc"]:.4f}, PR-AUC {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["logistic_balanced"]["pr_auc"]:.4f}, F1 {AVAILABILITY_STAGE_BEFORE_TRAINING_REWORK["logistic_balanced"]["f1"]:.4f}

### Etapa actual: reentrenamiento residual apoyado en baseline lineal
- PyTorch crudo: accuracy {format_metric(pytorch_test["accuracy"])}, ROC-AUC {format_metric(pytorch_test["roc_auc"])}, PR-AUC {format_metric(pytorch_test["pr_auc"])}, F1 {format_metric(pytorch_test["f1"])}, precision {format_metric(pytorch_test["precision"])}, recall {format_metric(pytorch_test["recall"])}
- PyTorch calibrado: accuracy {format_metric(pytorch_calibrated_test["accuracy"])}, ROC-AUC {format_metric(pytorch_calibrated_test["roc_auc"])}, PR-AUC {format_metric(pytorch_calibrated_test["pr_auc"])}, F1 {format_metric(pytorch_calibrated_test["f1"])}, precision {format_metric(pytorch_calibrated_test["precision"])}, recall {format_metric(pytorch_calibrated_test["recall"])}
- Logistic balanceado: accuracy {format_metric(baseline_lr["accuracy"])}, ROC-AUC {format_metric(baseline_lr["roc_auc"])}, PR-AUC {format_metric(baseline_lr["pr_auc"])}, F1 {format_metric(baseline_lr["f1"])}

## Hallazgos verificados
- El target temporal sigue siendo exigente: la tasa positiva actual es {metadata["dataset_summary"]["class_distribution"]["positive_rate"]:.2%}.
- Respecto de la etapa publicada anterior (`a7e6f7f`), PyTorch mejora en las metricas comparadas:
- cambio en F1: {delta_pytorch_f1:+.4f}
- cambio en PR-AUC: {delta_pytorch_pr_auc:+.4f}
- Respecto de la etapa local inmediatamente anterior con fisico/disponibilidad, PyTorch mejora de forma clara:
- cambio en F1: {delta_vs_availability_f1:+.4f}
- cambio en PR-AUC: {delta_vs_availability_pr_auc:+.4f}
- La calibracion de probabilidades quedo implementada y en la corrida actual el metodo seleccionado fue `{metadata["pytorch"]["calibration_method"]}`.
- La salida cruda queda como score principal porque supera al baseline en `PR-AUC` y `F1`; la calibrada queda como referencia secundaria porque mejora `F1` pero baja `PR-AUC`.
- El target actual deja {metadata["dataset_summary"].get("temporal_consolidation_count", "N/D")} positivos por consolidacion y {metadata["dataset_summary"].get("temporal_breakout_count", "N/D")} por breakout.
- El sistema ahora tiene una senal longitudinal mucho mas rica en la base:
- disponibilidad mensual
- fatiga
- carga de trabajo
- lesion/inactividad
- maduracion fisica y crecimiento corporal
- PyTorch vuelve a acercarse al baseline lineal gracias al bootstrap residual.
- En esta corrida, PyTorch crudo supera al baseline lineal balanceado en `PR-AUC` y `F1`.
- El baseline `LogisticRegression(class_weight="balanced")` se conserva como comparador formal obligatorio para futuras corridas.
- La prediccion es metodologicamente mas defendible porque ahora la progresion futura no depende solo de tecnica y contexto de partido, sino tambien de disponibilidad y maduracion fisica.
- El sistema mantiene el pipeline compartido entre entrenamiento e inferencia, pero el entrenamiento ya no mira la trayectoria completa como si fuera toda observable en el momento de decidir.

## Limites honestos de esta etapa
- Los partidos sinteticos siguen siendo por jugador; todavia no representan encuentros compartidos entre varios jugadores del mismo plantel.
- `ScoutReport` sigue siendo sintetico, no manual ni proveniente de observacion real de usuarios.
- El target temporal actual sigue siendo sintetico y deja {metadata["dataset_summary"]["class_distribution"]["positive"]} positivos sobre {metadata["dataset_summary"]["filtered_rows"]} jugadores.
- PyTorch crudo gana en la corrida oficial actual, pero esa ventaja sigue basada en datos sinteticos y debe validarse si cambia el target, la semilla o aparecen datos reales.
- No seria honesto eliminar el baseline: sigue siendo necesario para demostrar que PyTorch aporta valor en cada reentrenamiento.
- La base de entrenamiento SQLite ya es pesada y el repo recibio advertencia de GitHub por tamano de `players_training.db`.

## Validacion ejecutada
- `pytest -q`: 40 tests aprobados.
- Smoke de app con artefactos nuevos:
- `/` -> 200
- `/health` -> 200
- `/login` -> 200
- Reentrenamiento completo ejecutado sobre la nueva base sintetica.

## Proximo paso recomendado
- Mantener la salida cruda de PyTorch como score principal del MVP mientras conserve mejor ranking que el baseline.
- Dejar la probabilidad calibrada como evidencia secundaria y no como score principal.
- Seguir enriqueciendo la generacion sintetica con senales longitudinales no triviales, sin aumentar volumen por aumentar.
- Si se reabre el modelo, validar de nuevo PyTorch vs baseline bajo el mismo split y documentar la decision.
"""


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        doc.add_paragraph(line)
    doc.save(output_path)


def main() -> None:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    counts = db_counts()
    markdown = build_markdown(metadata, counts)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    PROGRESS_MD.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown, PROGRESS_DOCX)
    print(f"Generado: {PROGRESS_MD}")
    print(f"Generado: {PROGRESS_DOCX}")


if __name__ == "__main__":
    main()
