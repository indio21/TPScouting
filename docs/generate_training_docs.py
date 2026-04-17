from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
SCOUTING_APP_DIR = REPO_ROOT / "scouting_app"
DOCS_DIR = REPO_ROOT / "docs"
METADATA_PATH = SCOUTING_APP_DIR / "training_metadata.json"
EVIDENCE_MD = DOCS_DIR / "model_training_evidence.md"
EVIDENCE_DOCX = DOCS_DIR / "model_training_evidence.docx"
PLAN_MD = DOCS_DIR / "model_training_plan.md"
PLAN_DOCX = DOCS_DIR / "model_training_plan.docx"

PRIOR_DIAGNOSIS = {
    "dataset_rows": 20000,
    "positives": 3038,
    "negatives": 16962,
    "positive_rate": 0.1519,
    "age_min": 12,
    "age_max": 28,
    "pytorch": {
        "accuracy": 0.8480,
        "roc_auc": 0.2123,
        "pr_auc": 0.0915,
        "f1": 0.0,
        "precision": 0.0,
        "recall": 0.0,
        "confusion_matrix": [[3392, 0], [608, 0]],
    },
    "logistic_default": {
        "roc_auc": 0.9072,
        "pr_auc": 0.6258,
        "f1": 0.5132,
    },
    "logistic_balanced": {
        "roc_auc": 0.9074,
        "pr_auc": 0.6261,
        "f1": 0.5777,
        "recall": 0.8257,
    },
    "logistic_balanced_12_18": {
        "roc_auc": 0.9249,
        "pr_auc": 0.6531,
        "f1": 0.5932,
    },
    "avg_skill": {
        "roc_auc": 0.9162,
        "pr_auc": 0.6434,
    },
}


def git_branch() -> str:
    try:
        return (
            subprocess.check_output(
                ["git", "-C", str(REPO_ROOT), "rev-parse", "--abbrev-ref", "HEAD"],
                text=True,
            )
            .strip()
        )
    except Exception:
        return "desconocida"


def format_metric(value) -> str:
    if value in ("", None):
        return "N/D"
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def build_evidence_markdown(metadata: dict) -> str:
    dataset_summary = metadata["dataset_summary"]
    pytorch_test = metadata["pytorch"]["test"]
    pytorch_val = metadata["pytorch"]["validation"]
    baseline_lr = metadata["baselines"]["logistic_regression_balanced"]["test"]
    baseline_avg = metadata["baselines"]["avg_skill_score"]["test"]

    return f"""# Evidencia tecnica del entrenamiento

## Estado actual de la rama
- Rama analizada: `{git_branch()}`
- Cambio estructural ya incorporado: pipeline compartido de preprocesamiento con `pandas`, `ColumnTransformer`, `SimpleImputer`, `MinMaxScaler`, `OneHotEncoder` y persistencia en `preprocessor.joblib`.
- Artefactos actuales del entrenamiento: `model.pt`, `preprocessor.joblib`, `training_metadata.json` y `experiments.csv`.
- El objetivo del producto se mantiene acotado a scouting juvenil de 12 a 18 anos para clubes formativos.

## Diagnostico previo verificado antes de endurecer el entrenamiento
- Dataset original analizado: {PRIOR_DIAGNOSIS["dataset_rows"]} jugadores sinteticos.
- Distribucion de clases original: {PRIOR_DIAGNOSIS["positives"]} positivos y {PRIOR_DIAGNOSIS["negatives"]} negativos.
- Tasa positiva original: {PRIOR_DIAGNOSIS["positive_rate"]:.2%}.
- Rango etario original del dataset: {PRIOR_DIAGNOSIS["age_min"]}-{PRIOR_DIAGNOSIS["age_max"]}.
- Resultado real de la MLP anterior: accuracy {PRIOR_DIAGNOSIS["pytorch"]["accuracy"]:.4f}, ROC-AUC {PRIOR_DIAGNOSIS["pytorch"]["roc_auc"]:.4f}, PR-AUC {PRIOR_DIAGNOSIS["pytorch"]["pr_auc"]:.4f}, F1 {PRIOR_DIAGNOSIS["pytorch"]["f1"]:.4f}.
- Interpretacion honesta del diagnostico previo: la red colapsaba a clase negativa y la accuracy era enganosa.
- Baseline `LogisticRegression` anterior: ROC-AUC {PRIOR_DIAGNOSIS["logistic_default"]["roc_auc"]:.4f}, PR-AUC {PRIOR_DIAGNOSIS["logistic_default"]["pr_auc"]:.4f}, F1 {PRIOR_DIAGNOSIS["logistic_default"]["f1"]:.4f}.
- Baseline `LogisticRegression(class_weight="balanced")` anterior: ROC-AUC {PRIOR_DIAGNOSIS["logistic_balanced"]["roc_auc"]:.4f}, PR-AUC {PRIOR_DIAGNOSIS["logistic_balanced"]["pr_auc"]:.4f}, F1 {PRIOR_DIAGNOSIS["logistic_balanced"]["f1"]:.4f}, recall {PRIOR_DIAGNOSIS["logistic_balanced"]["recall"]:.4f}.
- Baseline simple por promedio de atributos anterior: ROC-AUC {PRIOR_DIAGNOSIS["avg_skill"]["roc_auc"]:.4f}, PR-AUC {PRIOR_DIAGNOSIS["avg_skill"]["pr_auc"]:.4f}.

## Cambios implementados en esta iteracion
- Se reemplazo `BCELoss` por `BCEWithLogitsLoss`.
- Se elimino la `Sigmoid` final de `PlayerNet` y ahora se trabaja con logits; la probabilidad se recupera con `torch.sigmoid(...)` solo en evaluacion e inferencia.
- Se incorporo manejo explicito del desbalance con `pos_weight`.
- Se cambio el split a `train / validation / test` con seleccion del threshold en validacion.
- Se agrego early stopping sobre `PR-AUC` con desempate por `F1` positiva.
- El entrenamiento ahora se alinea por defecto al alcance real del MVP: edades 12-18.
- La generacion sintetica de `potential_label` ahora usa score ponderado por posicion, ajuste etario, componente mental y ruido controlado.
- Se formalizo `LogisticRegression(class_weight="balanced")` como baseline obligatorio de comparacion.
- Se guarda metadata completa de entrenamiento en `training_metadata.json`.
- Se agregaron features historicas agregadas al pipeline:
- cantidad de registros por jugador
- promedio historico de `final_score`
- promedio historico de `pass_accuracy`
- ultimo `final_score` registrado
- La generacion sintetica ahora crea `PlayerStat` para que esas features existan tambien en la base de entrenamiento.
- `PlayerAttributeHistory` ahora entra al pipeline con senales longitudinales como:
- mejora media en 90, 180 y 365 dias
- pendiente de crecimiento
- volatilidad del progreso
- gap entre la ficha actual y la trayectoria reciente
- La generacion sintetica ahora crea entre 6 y 12 snapshots tecnicos por jugador y deriva `PlayerStat` desde esa evolucion.

## Resultado actual del entrenamiento mejorado
- Fecha de corrida registrada: `{metadata["timestamp"]}`
- Dataset actual: {dataset_summary["filtered_rows"]} jugadores dentro del rango {dataset_summary["age_range_filtered"]["min"]}-{dataset_summary["age_range_filtered"]["max"]}.
- Distribucion actual de clases: {dataset_summary["class_distribution"]["positive"]} positivos y {dataset_summary["class_distribution"]["negative"]} negativos.
- Tasa positiva actual: {dataset_summary["class_distribution"]["positive_rate"]:.2%}.
- Split efectivo: train {metadata["dataset"]["train_size"]}, validation {metadata["dataset"]["validation_size"]}, test {metadata["dataset"]["test_size"]}.
- `pos_weight` utilizado: {metadata["config"]["pos_weight"]:.4f}.
- Early stopping: mejor epoca {metadata["pytorch"]["best_epoch"]} y threshold elegido {metadata["pytorch"]["selected_threshold"]:.3f}.

## Metricas del modelo PyTorch actual
- Validacion: accuracy {format_metric(pytorch_val["accuracy"])}, ROC-AUC {format_metric(pytorch_val["roc_auc"])}, PR-AUC {format_metric(pytorch_val["pr_auc"])}, F1 {format_metric(pytorch_val["f1"])}, precision {format_metric(pytorch_val["precision"])}, recall {format_metric(pytorch_val["recall"])}.
- Test: accuracy {format_metric(pytorch_test["accuracy"])}, ROC-AUC {format_metric(pytorch_test["roc_auc"])}, PR-AUC {format_metric(pytorch_test["pr_auc"])}, F1 {format_metric(pytorch_test["f1"])}, precision {format_metric(pytorch_test["precision"])}, recall {format_metric(pytorch_test["recall"])}.
- Matriz de confusion PyTorch en test: {pytorch_test["confusion_matrix"]}.

## Baselines actuales bajo el mismo split y preprocesamiento
- `LogisticRegression(class_weight="balanced")`: accuracy {format_metric(baseline_lr["accuracy"])}, ROC-AUC {format_metric(baseline_lr["roc_auc"])}, PR-AUC {format_metric(baseline_lr["pr_auc"])}, F1 {format_metric(baseline_lr["f1"])}, precision {format_metric(baseline_lr["precision"])}, recall {format_metric(baseline_lr["recall"])}.
- Baseline simple por promedio de atributos: accuracy {format_metric(baseline_avg["accuracy"])}, ROC-AUC {format_metric(baseline_avg["roc_auc"])}, PR-AUC {format_metric(baseline_avg["pr_auc"])}, F1 {format_metric(baseline_avg["f1"])}.

## Hallazgos verificados
- El nuevo preprocesamiento compartido con `pandas` y `scikit-learn` quedo implementado y funcionando tanto en entrenamiento como en inferencia.
- La MLP actual ya no colapsa a todo negativo: paso de F1 0.0000 a F1 {format_metric(pytorch_test["f1"])} y de PR-AUC 0.0915 a PR-AUC {format_metric(pytorch_test["pr_auc"])}.
- El entrenamiento ya no usa solo foto fija: aprende con rendimiento historico y con evolucion tecnica de `PlayerAttributeHistory`.
- La alineacion del dataset a 12-18, el entrenamiento endurecido y las features longitudinales mejoraron fuerte la defendibilidad metodologica respecto al diagnostico previo.
- Aun asi, el baseline `LogisticRegression(class_weight="balanced")` sigue superando a la MLP en ROC-AUC, PR-AUC y F1.
- El baseline simple por promedio de atributos ya no explica bien el target frente al nuevo pipeline, lo que indica que la etiqueta sintetica quedo menos trivial que antes.
- La senal del dataset existe, pero la red PyTorch todavia no demuestra una ventaja clara sobre el baseline lineal balanceado.

## Limites que todavia no estan resueltos
- Todavia falta contexto de partido explicito (`Match` y participacion del jugador por partido).
- Todavia no hay reportes cualitativos de scout en la base de entrenamiento.
- El target sigue siendo `potential_label` binario y no una meta temporal de progresion.
- No se implemento calibracion de probabilidades.
- La evidencia actual sigue basada en datos sinteticos; no hay una validacion externa con datos reales.
- La MLP mejoro, pero todavia no justifica por rendimiento reemplazar al baseline lineal como referencia formal.

## Pruebas ejecutadas
- `pytest -q`: 33 tests aprobados.
- Cobertura nueva o reforzada:
- persistencia del `preprocessor.joblib`
- consistencia entre inferencia individual y batch
- split `train / validation / test` y metadata de threshold
- filtro de entrenamiento acotado a 12-18
- sensibilidad de la etiqueta sintetica a edad y posicion
- merge de features historicas en el dataset de entrenamiento
- inclusion de features historicas en la inferencia de la app
- features longitudinales de `PlayerAttributeHistory`
- smoke del pipeline completo de entrenamiento

## Procedimiento reproducible
- Regenerar dataset de entrenamiento:
- `python scouting_app/generate_data.py --num-players 20000 --db-url sqlite:///players_training.db --seed 42`
- Reentrenar artefactos:
- `python scouting_app/train_model.py --db-url sqlite:///players_training.db --model-out scouting_app/model.pt --preprocessor-out scouting_app/preprocessor.joblib --metadata-out scouting_app/training_metadata.json --epochs 30 --lr 1e-3 --patience 8`
- Ejecutar tests:
- `python -m pytest -q`
"""


def build_plan_markdown(metadata: dict) -> str:
    baseline_lr = metadata["baselines"]["logistic_regression_balanced"]["test"]
    pytorch_test = metadata["pytorch"]["test"]
    return f"""# Plan tecnico de entrenamiento

## Resumen
- Este plan queda versionado para la rama `training`.
- PyTorch sigue siendo el modelo principal del MVP.
- `LogisticRegression(class_weight="balanced")` queda formalizada como baseline obligatorio de comparacion.
- La documentacion tecnica y la evidencia se guardan separadas del documento final de tesis.

## Etapas implementadas en esta iteracion
### Etapa 1 completada: endurecer entrenamiento PyTorch
- `BCEWithLogitsLoss` en lugar de `BCELoss`.
- Modelo sin `Sigmoid` final y trabajo con logits.
- `pos_weight` aplicado por defecto para desbalance.
- Split `train / validation / test`.
- Seleccion de threshold por validacion.
- Early stopping sobre `PR-AUC` con desempate por `F1`.
- Resultado: PyTorch dejo de colapsar y alcanzo F1 {format_metric(pytorch_test["f1"])} en test.

### Etapa 2 completada: alinear dataset al alcance real 12-18
- Generacion sintetica por defecto restringida a 12-18.
- Entrenamiento por defecto restringido a 12-18.
- Resultado: el pipeline ahora refleja el alcance real del producto.

### Etapa 3 completada: hacer mas realista `potential_label`
- Score ponderado por posicion.
- Ajuste etario juvenil.
- Mayor incidencia de `determination`, `technique` y `vision`.
- Ruido controlado en lugar de regla casi lineal por promedio simple.

### Etapa 4 completada: formalizar baseline y comparacion
- Baseline obligatorio: `LogisticRegression(class_weight="balanced")`.
- Comparacion bajo el mismo split y el mismo preprocesamiento.
- Persistencia de metadata en `training_metadata.json`.
- Estado actual: el baseline sigue mejor que PyTorch en test con F1 {format_metric(baseline_lr["f1"])} vs {format_metric(pytorch_test["f1"])}.

### Etapa 5 completada en nivel MVP
- Se persisten threshold, metricas, tamanos de split, seed y configuracion.
- La evidencia tecnica puede regenerarse desde los artefactos del repo.

### Etapa 6 completada: crecimiento con features historicas
- Se agregaron features historicas agregadas con `pandas`.
- Se sintetizo historial de `PlayerStat` en la base de entrenamiento.
- Se integraron features de `PlayerAttributeHistory` al entrenamiento e inferencia.
- La base de entrenamiento ahora representa trayectoria tecnica del jugador, no solo foto fija.
- Resultado actual: PyTorch queda en F1 {format_metric(pytorch_test["f1"])} y PR-AUC {format_metric(pytorch_test["pr_auc"])}.
- Aun asi, el baseline lineal balanceado sigue siendo superior.

## Siguiente iteracion recomendada
### Etapa 7 recomendada: contexto de partido y target temporal
- Agregar una tabla `Match` con contexto minimo del partido.
- Agregar una tabla de participacion del jugador por partido.
- Redefinir el target hacia una meta temporal de progresion, no solo `potential_label`.
- Volver a comparar PyTorch contra el baseline lineal bajo el mismo split.

## Criterios de aceptacion para la siguiente iteracion
- La siguiente iteracion debe mejorar o estabilizar al menos una de estas metricas de PyTorch en test sin degradar claramente las demas:
- PR-AUC
- F1 positiva
- Recall positiva
- Y ademas debe volver mas explicable la prediccion desde el punto de vista futbolistico.
- La comparacion debe seguir quedando trazable en `training_metadata.json`.
- Si PyTorch no supera al baseline tras esa iteracion, el baseline debe quedar reconocido como referencia principal de rendimiento.

## Pruebas de validacion vigentes
- Tests del preprocesador compartido.
- Tests de consistencia entre inferencia individual y batch.
- Tests del split `train / validation / test`.
- Tests del threshold seleccionado y su persistencia.
- Tests de entrenamiento limitado a 12-18.
- Tests de generacion sintetica sensibles a edad y posicion.
- Tests de merge de features historicas.
- Tests de features longitudinales de `PlayerAttributeHistory`.
- Smoke del pipeline completo con artefactos reales.

## Decisiones fijas
- No se toca todavia el documento final de tesis.
- Todo este trabajo se mantiene en la rama `training` hasta que se decida mergearlo.
- El producto sigue acotado a scouting juvenil 12-18 para clubes formativos sin grandes presupuestos.
"""


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        if line[:3].isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line[3:], style="List Number")
            continue
        doc.add_paragraph(line)
    doc.save(output_path)


def main() -> None:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    evidence_markdown = build_evidence_markdown(metadata)
    plan_markdown = build_plan_markdown(metadata)

    EVIDENCE_MD.write_text(evidence_markdown, encoding="utf-8")
    PLAN_MD.write_text(plan_markdown, encoding="utf-8")

    markdown_to_docx(evidence_markdown, EVIDENCE_DOCX)
    markdown_to_docx(plan_markdown, PLAN_DOCX)

    print(f"Generado: {EVIDENCE_MD}")
    print(f"Generado: {EVIDENCE_DOCX}")
    print(f"Generado: {PLAN_MD}")
    print(f"Generado: {PLAN_DOCX}")


if __name__ == "__main__":
    main()
